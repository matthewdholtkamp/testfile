import csv
import json
import os
import re
import shutil
import ast
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
STATE_DIR = REPO_ROOT / 'outputs' / 'state'
MANUSCRIPT_ROOT = REPO_ROOT / 'outputs' / 'manuscripts'
MANUSCRIPT_DRAFTS_ROOT = REPO_ROOT / 'Manuscript Drafts' / 'Generated'
READY_FOR_METADATA_DRAFTS_ROOT = REPO_ROOT / 'Manuscript Drafts' / 'Ready for Metadata Only'
OUTPUT_ROOT = REPO_ROOT / 'output'
JOURNAL_REGISTRY_PATH = REPO_ROOT / 'config' / 'manuscript_journal_registry.json'
JOURNAL_REQUIREMENTS_DIR = REPO_ROOT / 'config' / 'journal_requirements'
MANUSCRIPT_QUEUE_STATE_PATH = STATE_DIR / 'manuscript_queue.json'
PUBLICATION_TRACKER_STATE_PATH = STATE_DIR / 'publication_tracker.json'
READY_FOR_CODEX_FILENAME = 'ready_for_codex_draft.json'
GENERATED_DRAFT_INDEX_FILENAME = 'generated_draft_index.json'
GENERATED_DRAFT_MANIFEST_FILENAME = 'draft_manifest.json'
READY_FOR_METADATA_INDEX_FILENAME = 'ready_for_metadata_only_index.json'
ACTIVE_LANE_LIMIT = 2
MANUSCRIPT_CANDIDATE_EVALUATION_LIMIT = 12
PROCESS_ENGINE_DOC_PATH = REPO_ROOT / 'docs' / 'process-engine' / 'process_engine.json'

QUEUE_STATUSES = [
    'awaiting your approval',
    'approved for submission',
    'submitted',
    'under review',
    'accepted',
    'published',
    'rejected / retargeting',
]
SUBMISSION_TRACKER_EXIT_STATUSES = {'submitted', 'under review', 'accepted', 'published', 'rejected / retargeting'}
ACTIVE_APPROVAL_STATUSES = {'awaiting your approval', 'approved for submission'}
TASK_STATES = ['proposed', 'queued', 'running', 'satisfied', 'blocked', 'stale', 'reopened']
CONNECTOR_SOURCE_LABELS = {
    'open_targets': 'Open Targets',
    'clinicaltrials_gov': 'ClinicalTrials.gov',
    'biorxiv_medrxiv': 'bioRxiv / medRxiv',
    'tenx_genomics': '10x Genomics',
    'chembl': 'ChEMBL',
}
CONNECTOR_MATCH_STOPWORDS = {
    'most',
    'informative',
    'biomarker',
    'panel',
    'packet',
    'candidate',
    'paper',
    'manuscript',
    'path',
    'lead',
}


def normalize(value):
    if value is None:
        return ''
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def slugify(value):
    value = normalize(value).lower()
    if not value:
        return 'unnamed'
    value = re.sub(r'[^a-z0-9]+', '_', value)
    return value.strip('_') or 'unnamed'


def canonical_candidate_id(value):
    return normalize(value).replace(' ', '-').replace(':', '-').replace('/', '-').lower()


def pack_key(value):
    return slugify(normalize(value).replace(':', '-').replace('/', '-'))


def read_json(path, default=None):
    path = Path(path)
    if not path.exists():
        return deepcopy(default)
    with path.open() as handle:
        return json.load(handle)


def write_json(path, payload):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open('w') as handle:
        json.dump(payload, handle, indent=2)
        handle.write('\n')


def write_text(path, text):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text)


def read_csv(path):
    path = Path(path)
    if not path.exists():
        return []
    with path.open(newline='', encoding='utf-8') as handle:
        return list(csv.DictReader(handle))


def iso_now():
    return datetime.now(timezone.utc).isoformat()


def pretty_label(value):
    value = normalize(value)
    if not value:
        return 'Not specified'
    replacements = {
        'bbb': 'BBB',
        'ros': 'ROS',
        'aqp4': 'AQP4',
        'nlrp3': 'NLRP3',
        'prkn': 'PRKN',
        'dti': 'DTI',
        'cbf': 'CBF',
        'nfl': 'NfL',
        'gfap': 'GFAP',
        'ocln': 'OCLN',
        'cldn5': 'CLDN5',
        'tjp1': 'TJP1',
        'il': 'IL',
    }
    pieces = []
    for part in value.replace('-', '_').split('_'):
        lower = part.lower()
        if lower in replacements:
            pieces.append(replacements[lower])
        elif lower.isdigit():
            pieces.append(lower)
        else:
            pieces.append(lower.capitalize())
    return ' '.join(pieces)


def summarize_list(values, limit=4):
    clean = [pretty_label(item) for item in (values or []) if normalize(item)]
    if not clean:
        return 'not specified'
    if len(clean) <= limit:
        return ', '.join(clean)
    return ', '.join(clean[:limit]) + f', and {len(clean) - limit} more'


def connector_source_label(source):
    source = normalize(source)
    return CONNECTOR_SOURCE_LABELS.get(source, pretty_label(source))


def latest_matching_path(patterns):
    candidates = []
    for pattern in patterns:
        candidates.extend(REPO_ROOT.glob(pattern))
    if not candidates:
        return None
    return sorted(candidates)[-1]


def relative_path_or_blank(path):
    if not path:
        return ''
    try:
        return str(Path(path).resolve().relative_to(REPO_ROOT))
    except ValueError:
        return str(path)


def absolute_path_or_blank(path):
    if not path:
        return ''
    return str(Path(path).resolve())


def tokenize_for_connector_match(*values):
    tokens = set()
    for value in values:
        for token in re.findall(r'[a-z0-9]+', normalize(value).lower()):
            if len(token) < 3 or token in CONNECTOR_MATCH_STOPWORDS:
                continue
            tokens.add(token)
    return tokens


def connector_match_score(candidate, value):
    candidate_tokens = tokenize_for_connector_match(
        candidate.get('theme_key'),
        candidate.get('theme_label'),
        candidate.get('title'),
    )
    value_tokens = tokenize_for_connector_match(value)
    return len(candidate_tokens.intersection(value_tokens))


def count_csv_rows(path):
    if not path or not Path(path).exists():
        return 0
    with Path(path).open(newline='', encoding='utf-8') as handle:
        return sum(1 for _ in csv.DictReader(handle))


def sorted_counter_rows(counter):
    return [
        {
            'key': key,
            'label': connector_source_label(key),
            'count': counter[key],
        }
        for key in sorted(counter)
    ]


def build_connector_context():
    enrichment_path = latest_matching_path([
        'reports/connector_enrichment_curated/connector_enrichment_curated_*.csv',
        'reports/connector_enrichment/connector_enrichment_records_*.csv',
        'reports/connector_enrichment/connector_enrichment_*.csv',
    ])
    manifest_path = latest_matching_path([
        'reports/connector_candidate_manifest/connector_candidate_manifest_*.csv',
    ])
    tenx_template_path = latest_matching_path([
        'reports/connector_candidate_manifest/templates/tenx_genomics_import_template_*.csv',
        'local_connector_inputs/templates/tenx_genomics_import_template_*.csv',
    ])

    enrichment_rows = read_csv(enrichment_path) if enrichment_path else []
    manifest_rows = read_csv(manifest_path) if manifest_path else []
    tenx_template_rows = read_csv(tenx_template_path) if tenx_template_path else []
    latest_retrieved_at = max(
        (normalize(row.get('retrieved_at')) for row in enrichment_rows if normalize(row.get('retrieved_at'))),
        default='',
    )
    return {
        'enrichment_path': enrichment_path,
        'enrichment_rows': enrichment_rows,
        'manifest_path': manifest_path,
        'manifest_rows': manifest_rows,
        'tenx_template_path': tenx_template_path,
        'tenx_template_rows': tenx_template_rows,
        'latest_retrieved_at': latest_retrieved_at,
    }


def build_connector_dashboard_payload(connector_context=None):
    connector_context = connector_context or build_connector_context()
    enrichment_rows = connector_context.get('enrichment_rows') or []
    by_source = {}
    by_tier = {}
    tenx_rows = 0
    for row in enrichment_rows:
        source = normalize(row.get('connector_source'))
        tier = normalize(row.get('evidence_tier'))
        if source:
            by_source[source] = by_source.get(source, 0) + 1
        if tier:
            by_tier[tier] = by_tier.get(tier, 0) + 1
        if source == 'tenx_genomics' or tier == 'genomics_expression':
            tenx_rows += 1
    tenx_template_rows = connector_context.get('tenx_template_rows') or []
    tenx_state = 'no_10x_template_prepared'
    if tenx_rows:
        tenx_state = 'real_10x_evidence_attached'
    elif tenx_template_rows:
        tenx_state = 'template_only'
    return {
        'latest_enrichment_path': relative_path_or_blank(connector_context.get('enrichment_path')),
        'latest_enrichment_at': connector_context.get('latest_retrieved_at') or '',
        'row_count': len(enrichment_rows),
        'by_source': sorted_counter_rows(by_source),
        'by_evidence_tier': [
            {
                'key': key,
                'label': pretty_label(key),
                'count': by_tier[key],
            }
            for key in sorted(by_tier)
        ],
        'tenx': {
            'state': tenx_state,
            'imported_row_count': tenx_rows,
            'template_row_count': len(tenx_template_rows),
            'template_path': relative_path_or_blank(connector_context.get('tenx_template_path')),
        },
    }


def split_notes(value):
    if isinstance(value, list):
        raw_parts = []
        for item in value:
            raw_parts.extend(split_notes(item))
        return ordered_unique(raw_parts)
    text = normalize(value)
    if not text:
        return []
    if text.startswith('[') and text.endswith(']'):
        try:
            parsed = ast.literal_eval(text)
        except (ValueError, SyntaxError):
            parsed = None
        if isinstance(parsed, list):
            return split_notes(parsed)
    bits = [normalize(part) for part in re.split(r'\|\||;|\n', text) if normalize(part)]
    bits = [part for part in bits if part.lower() not in {'none', 'none_detected', 'not specified', 'n/a'}]
    return bits


def default_publication_tracker():
    return {
        'updated_at': '',
        'entries': [],
        'valid_statuses': QUEUE_STATUSES,
        'notes': 'Manual by operator for now. Move manuscripts here once they reach human approval or journal-facing states.',
    }


def load_publication_tracker(path=PUBLICATION_TRACKER_STATE_PATH):
    tracker = read_json(path, default_publication_tracker()) or {}
    merged = default_publication_tracker()
    merged.update(tracker)
    merged['entries'] = list(merged.get('entries') or [])
    return merged


def load_journal_registry(path=JOURNAL_REGISTRY_PATH):
    data = read_json(path, {'journals': []}) or {'journals': []}
    data['journals'] = list(data.get('journals') or [])
    return data


def family_preference(family_id):
    return {
        'most_informative_biomarker_panel': 4,
        'strongest_causal_bridge': 3,
        'best_intervention_leverage_point': 2,
        'highest_value_next_task': 1,
        'weakest_evidence_hinge': 0,
    }.get(normalize(family_id), 0)


def support_weight(support_status):
    return {'supported': 2, 'provisional': 1, 'weak': 0}.get(normalize(support_status), 0)


def theme_key(row):
    display = normalize(row.get('display_name'))
    if display:
        return slugify(display)
    title = normalize(row.get('title'))
    title = re.sub(r'\s+(biomarker panel|discriminator panel)$', '', title, flags=re.IGNORECASE)
    return slugify(title)


def journal_requirements_path(journal_id):
    journal_id = normalize(journal_id)
    if not journal_id:
        return None
    return JOURNAL_REQUIREMENTS_DIR / f'{journal_id}.json'


def load_journal_requirements(journal_id):
    path = journal_requirements_path(journal_id)
    if not path or not path.exists():
        return {}
    return read_json(path, {}) or {}


def latest_corpus_manifest_path():
    candidates = sorted(OUTPUT_ROOT.glob('pubmed_tbi_manifest_*.csv'))
    return candidates[-1] if candidates else None


def parse_pmid_list(value):
    if isinstance(value, list):
        raw_values = value
    else:
        raw_values = re.split(r'[;,\|\n]+', normalize(value))
    seen = set()
    pmids = []
    for item in raw_values:
        text = normalize(item)
        if not text:
            continue
        match = re.search(r'(\d{6,9})', text)
        if not match:
            continue
        pmid = match.group(1)
        if pmid not in seen:
            seen.add(pmid)
            pmids.append(pmid)
    return pmids


def source_quality_tier_from_rank(rank_value):
    rank = normalize(rank_value)
    if rank in {'2', '3', '4', '5'}:
        return 'full_text_like'
    if rank == '1':
        return 'abstract_only'
    return 'unknown'


def parse_source_markdown_metadata(path):
    path = Path(path)
    if not path.exists():
        return {}
    relative_path = str(path.relative_to(REPO_ROOT))
    if os.environ.get('PHASE8_PARSE_SOURCE_MARKDOWN', '').strip() != '1':
        return {
            'source_markdown_path': relative_path,
        }

    text = path.read_text()

    def capture(label):
        match = re.search(rf'^\*\*{re.escape(label)}:\*\*\s*(.+)$', text, flags=re.MULTILINE)
        return normalize(match.group(1)) if match else ''

    abstract = ''
    abstract_match = re.search(r'## Abstract\s+(.+?)(?:\n## |\Z)', text, flags=re.DOTALL)
    if abstract_match:
        abstract = normalize(abstract_match.group(1))

    full_text_excerpt = ''
    full_text_match = re.search(r'## Full Text\s+(.+?)(?:\n## |\Z)', text, flags=re.DOTALL)
    if full_text_match:
        full_text_excerpt = normalize(full_text_match.group(1))[:2000]

    return {
        'title': normalize(text.splitlines()[0].lstrip('# ').strip()) if text else '',
        'authors': capture('Authors'),
        'journal': capture('Journal'),
        'publication_date': capture('Publication Date'),
        'pmid': capture('PMID'),
        'pmcid': capture('PMCID'),
        'doi': capture('DOI'),
        'pubmed_url': capture('PubMed URL'),
        'pmc_url': capture('PMC URL'),
        'extraction_source': capture('Extraction Source'),
        'extraction_rank': capture('Extraction Rank'),
        'abstract': abstract,
        'full_text_excerpt': full_text_excerpt,
        'source_markdown_path': relative_path,
    }


def load_corpus_reference_index():
    manifest_path = latest_corpus_manifest_path()
    manifest_rows = {}
    if manifest_path and manifest_path.exists():
        with manifest_path.open(newline='') as handle:
            for row in csv.DictReader(handle):
                pmid = normalize(row.get('PMID'))
                if pmid:
                    manifest_rows[pmid] = row

    source_paths = {}
    for path in OUTPUT_ROOT.glob('*_PMID*.md'):
        match = re.search(r'_PMID(\d+)\.md$', path.name)
        if match:
            source_paths[match.group(1)] = path

    return {
        'manifest_path': str(manifest_path.relative_to(REPO_ROOT)) if manifest_path else '',
        'manifest_rows': manifest_rows,
        'source_paths': source_paths,
        'parsed_markdown': {},
    }


def corpus_reference_record(pmid, corpus_index):
    pmid = normalize(pmid)
    if not pmid:
        return {}
    manifest_row = dict(corpus_index.get('manifest_rows', {}).get(pmid, {}))
    source_path = corpus_index.get('source_paths', {}).get(pmid)
    parsed_cache = corpus_index.setdefault('parsed_markdown', {})
    source_meta = parsed_cache.get(pmid)
    if source_meta is None:
        source_meta = parse_source_markdown_metadata(source_path) if source_path else {}
        parsed_cache[pmid] = source_meta

    extraction_rank = normalize(source_meta.get('extraction_rank') or manifest_row.get('extraction_rank'))
    extraction_source = normalize(source_meta.get('extraction_source') or manifest_row.get('extraction_source'))
    source_quality_tier = source_quality_tier_from_rank(extraction_rank)

    return {
        'pmid': pmid,
        'title': normalize(source_meta.get('title') or manifest_row.get('title')),
        'authors': normalize(source_meta.get('authors')),
        'journal': normalize(source_meta.get('journal')),
        'publication_date': normalize(source_meta.get('publication_date')),
        'doi': normalize(source_meta.get('doi') or manifest_row.get('DOI')),
        'pmcid': normalize(source_meta.get('pmcid') or manifest_row.get('PMCID')),
        'pubmed_url': normalize(source_meta.get('pubmed_url')),
        'pmc_url': normalize(source_meta.get('pmc_url')),
        'extraction_source': extraction_source,
        'extraction_rank': extraction_rank,
        'source_quality_tier': source_quality_tier,
        'manifest_source': corpus_index.get('manifest_path', ''),
        'source_markdown_path': normalize(source_meta.get('source_markdown_path')),
        'abstract': normalize(source_meta.get('abstract')),
        'full_text_excerpt': normalize(source_meta.get('full_text_excerpt')),
        'retrieval_mode': normalize(manifest_row.get('retrieval_mode')),
        'source_query': normalize(manifest_row.get('source_query')),
        'domain': normalize(manifest_row.get('domain')),
    }


def citation_text(record):
    pieces = []
    if normalize(record.get('authors')):
        pieces.append(record['authors'])
    if normalize(record.get('title')):
        pieces.append(record['title'])
    journal_and_year = ' '.join(part for part in [normalize(record.get('journal')), normalize(record.get('publication_date'))] if part)
    if journal_and_year:
        pieces.append(journal_and_year)
    if normalize(record.get('doi')):
        pieces.append(f"doi:{record['doi']}")
    pieces.append(f"PMID:{record['pmid']}")
    return '. '.join(piece.rstrip('.') for piece in pieces if piece).strip() + '.'


def load_process_rows(state):
    rows = list(state.get('process', {}).get('rows', []) or [])
    if rows:
        return rows
    if PROCESS_ENGINE_DOC_PATH.exists():
        data = read_json(PROCESS_ENGINE_DOC_PATH)
        return list(data.get('lanes', []) or [])
    return []


def evidence_chain_complete(row):
    return all([
        bool(row.get('linked_phase1_lane_ids')),
        bool(row.get('linked_phase2_transition_ids')),
        bool(row.get('linked_phase3_object_ids')),
        bool(row.get('linked_phase4_packet_ids')),
        bool(row.get('linked_phase5_endotype_ids')),
    ])


def build_phase_indexes(state):
    process_rows = load_process_rows(state)
    translational = {normalize(row.get('lane_id')): row for row in state.get('translational', {}).get('rows', []) or []}
    cohort = {normalize(row.get('packet_id') or row.get('endotype_id') or row.get('display_name')): row for row in state.get('cohort', {}).get('rows', []) or []}
    process = {normalize(row.get('lane_id') or row.get('display_name')): row for row in process_rows}
    transition = {normalize(row.get('transition_id') or row.get('display_name')): row for row in state.get('transition', {}).get('rows', []) or []}
    progression = {normalize(row.get('object_id') or row.get('display_name')): row for row in state.get('progression', {}).get('rows', []) or []}
    return {
        'translational': translational,
        'cohort': cohort,
        'process': process,
        'transition': transition,
        'progression': progression,
    }


def translation_maturity_for_row(row, indexes):
    maturities = []
    for lane_id in row.get('linked_phase4_packet_ids') or []:
        packet = indexes['translational'].get(normalize(lane_id)) or {}
        maturity = normalize(packet.get('translation_maturity') or packet.get('support_status'))
        if maturity:
            maturities.append(maturity)
    if 'actionable' in maturities:
        return 'actionable'
    if 'usable' in maturities:
        return 'usable'
    if 'bounded' in maturities:
        return 'bounded'
    return maturities[0] if maturities else 'not specified'


def candidate_modes(row):
    family = normalize(row.get('family_id'))
    title = normalize(row.get('title')).lower()
    modes = set()
    if family == 'most_informative_biomarker_panel':
        modes.add('mechanistic_biomarker')
    if family == 'strongest_causal_bridge':
        modes.add('mechanistic_bridge')
    if family == 'best_intervention_leverage_point':
        modes.add('translational_tbi')
    if family == 'highest_value_next_task':
        modes.add('methods_engine')
    if row.get('linked_phase5_endotype_ids'):
        modes.add('endotype_discriminator')
    if 'blast' in title or 'military' in title:
        modes.add('operational_tbi')
    if not modes:
        modes.add('clinical_synthesis')
    return sorted(modes)


def candidate_audiences(row):
    text = ' '.join([
        normalize(row.get('title')).lower(),
        normalize(row.get('display_name')).lower(),
        normalize(row.get('family_id')).lower(),
        normalize(row.get('blockers')).lower(),
    ])
    audiences = {'neurotrauma', 'mechanism'}
    if row.get('linked_phase5_endotype_ids'):
        audiences.add('cohort')
    if row.get('linked_phase4_packet_ids'):
        audiences.add('translational')
        audiences.add('biomarker')
    if 'blast' in text or 'military' in text:
        audiences.add('military')
    if 'critical' in text or 'acute severe' in text or 'vascular' in text:
        audiences.add('neurocritical')
    if 'rehab' in text or 'cognitive' in text or 'recovery' in text:
        audiences.add('rehabilitation')
    return sorted(audiences)


def candidate_article_types(row):
    family = normalize(row.get('family_id'))
    types = ['review']
    if family in {'most_informative_biomarker_panel', 'strongest_causal_bridge', 'best_intervention_leverage_point'}:
        types.extend(['original research', 'hypothesis and theory'])
    if family == 'highest_value_next_task':
        types.extend(['methodology', 'viewpoint'])
    if row.get('linked_phase5_endotype_ids'):
        types.append('brief report')
    ordered = []
    for value in types:
        if value not in ordered:
            ordered.append(value)
    return ordered


SEVERE_BLOCKER_TOKENS = [
    'no significant correlation',
    'contradict',
    'timing support is incomplete',
    'sparse',
    'thin',
    'still seeded',
    'abstract-only',
    'not enough',
]


def blocker_severity(row):
    notes = split_notes(row.get('blockers')) + split_notes(row.get('contradiction_notes')) + split_notes(row.get('evidence_gaps'))
    if not notes:
        return 'low'
    text_blob = ' '.join(note.lower() for note in notes)
    if any(token in text_blob for token in SEVERE_BLOCKER_TOKENS):
        return 'high'
    if len(notes) >= 2:
        return 'medium'
    return 'low'


def score_scientific_strength(row, indexes):
    score = 0
    score += support_weight(row.get('support_status'))
    if evidence_chain_complete(row):
        score += 1
    if len(row.get('linked_phase4_packet_ids') or []) and len(row.get('linked_phase5_endotype_ids') or []):
        score += 1
    severity = blocker_severity(row)
    if severity == 'high':
        score -= 1
    elif severity == 'medium':
        score -= 0.5
    maturity = translation_maturity_for_row(row, indexes)
    if maturity == 'bounded':
        score = min(score, 2)
    elif maturity in {'not specified', ''}:
        score = min(score, 2)
    return max(0, min(4, int(round(score))))


def journal_match_details(journal, candidate):
    audience_overlap = sorted(set(journal.get('audiences') or []).intersection(candidate['audiences']))
    mode_overlap = sorted(set(journal.get('manuscript_modes') or []).intersection(candidate['manuscript_modes']))
    type_overlap = sorted(set(journal.get('preferred_article_types') or []).intersection(candidate['article_type_candidates']))
    family_hit = normalize(candidate['family_id']) in set(journal.get('primary_fit_families') or [])
    score = 0
    score += min(3, len(audience_overlap))
    score += min(2, len(mode_overlap))
    score += 1 if type_overlap else 0
    score += 1 if family_hit else 0
    if journal.get('submission_tier') == 'stretch':
        score -= 1
    score = max(0, score)
    reasons = []
    if audience_overlap:
        reasons.append('Audience overlap: ' + ', '.join(pretty_label(item) for item in audience_overlap[:3]))
    if mode_overlap:
        reasons.append('Manuscript mode fit: ' + ', '.join(pretty_label(item) for item in mode_overlap[:3]))
    if type_overlap:
        reasons.append('Article-type overlap: ' + ', '.join(type_overlap[:2]))
    if family_hit:
        reasons.append('Family fit: journal often suits this manuscript family.')
    if not reasons:
        reasons.append('Weak direct fit; keep only as a distant backup.')
    return {
        'score': score,
        'audience_overlap': audience_overlap,
        'mode_overlap': mode_overlap,
        'article_type_overlap': type_overlap,
        'family_hit': family_hit,
        'reasons': reasons,
    }


def requirements_fit_details(journal, candidate):
    journal_id = normalize(journal.get('id'))
    path = journal_requirements_path(journal_id)
    snapshot = load_journal_requirements(journal_id)
    if not snapshot or not snapshot.get('verified'):
        unverified_reason = 'Requirements snapshot has not been captured yet. Keep this as a shortlist, not a verified fit.'
        if snapshot and not snapshot.get('verified'):
            unverified_reason = normalize(snapshot.get('notes')) or 'Requirements snapshot exists, but it is still marked unverified. Keep this as a shortlist, not a verified fit.'
        return {
            'checked': False,
            'score': 0,
            'article_type_match': False,
            'evidence_chain_match': False,
            'translation_match': False,
            'scientific_strength_match': False,
            'reasons': [unverified_reason],
            'requirements_path': str(path.relative_to(REPO_ROOT)) if path else '',
            'snapshot': snapshot or {},
        }

    allowed_article_types = snapshot.get('allowed_article_types') or snapshot.get('preferred_article_types') or []
    allowed_article_types = [normalize(item) for item in allowed_article_types if normalize(item)]
    candidate_article_types = [normalize(item) for item in candidate.get('article_type_candidates') or []]
    article_type_overlap = sorted(set(allowed_article_types).intersection(candidate_article_types))
    article_type_match = bool(article_type_overlap) if allowed_article_types else True

    evidence_expectations = snapshot.get('evidence_expectations') or {}
    requires_full_chain = bool(snapshot.get('requires_full_evidence_chain') or evidence_expectations.get('requires_full_evidence_chain'))
    evidence_chain_match = candidate.get('evidence_chain_complete', False) if requires_full_chain else True

    allows_bounded_translation = snapshot.get('allows_bounded_translation')
    if allows_bounded_translation is None:
        translation_match = True
    else:
        translation_match = bool(allows_bounded_translation) or normalize(candidate.get('translation_maturity')) != 'bounded'

    minimum_scientific_strength = int(snapshot.get('minimum_scientific_strength') or 0)
    scientific_strength_match = int(candidate.get('scientific_strength') or 0) >= minimum_scientific_strength

    reasons = []
    score = 0
    if article_type_match:
        score += 1
        reasons.append('Article type is compatible with the captured journal requirements.')
    else:
        reasons.append('Article type does not yet match the captured journal requirements.')
    if evidence_chain_match:
        score += 1
        reasons.append('Current evidence chain meets the journal requirements snapshot.')
    else:
        reasons.append('Current evidence chain is below the journal requirements snapshot.')
    if translation_match:
        score += 1
        reasons.append('Translation maturity is within the journal requirements bounds.')
    else:
        reasons.append('Bounded translation still conflicts with this journal requirements snapshot.')
    if scientific_strength_match:
        score += 1
        reasons.append('Scientific strength clears the minimum threshold recorded for this journal.')
    else:
        reasons.append('Scientific strength is still below the minimum threshold recorded for this journal.')

    return {
        'checked': True,
        'score': max(0, min(4, score)),
        'article_type_match': article_type_match,
        'evidence_chain_match': evidence_chain_match,
        'translation_match': translation_match,
        'scientific_strength_match': scientific_strength_match,
        'reasons': reasons,
        'requirements_path': str(path.relative_to(REPO_ROOT)) if path else '',
        'snapshot': snapshot,
    }


def dynamic_journal_pool(registry, candidate):
    journals = registry.get('journals') or []
    core = [journal for journal in journals if normalize(journal.get('registry_role')) == 'core']
    discovery = []
    for journal in journals:
        if normalize(journal.get('registry_role')) != 'discovery':
            continue
        overlap = set(journal.get('audiences') or []).intersection(candidate['audiences'])
        mode_overlap = set(journal.get('manuscript_modes') or []).intersection(candidate['manuscript_modes'])
        if overlap or mode_overlap:
            discovery.append(journal)
    return core, discovery


def choose_journal_targets(candidate, registry):
    core_pool, discovery_pool = dynamic_journal_pool(registry, candidate)
    seen = set()
    scored = []
    for journal in core_pool + discovery_pool:
        journal_id = normalize(journal.get('id'))
        if not journal_id or journal_id in seen:
            continue
        seen.add(journal_id)
        details = journal_match_details(journal, candidate)
        requirements = requirements_fit_details(journal, candidate)
        scored.append({
            'journal': journal,
            'score': details['score'],
            'reasons': details['reasons'],
            'article_type_overlap': details['article_type_overlap'],
            'requirements': requirements,
        })
    scored.sort(key=lambda item: (-item['score'], 0 if item['journal'].get('submission_tier') == 'realistic' else 1, normalize(item['journal'].get('name'))))

    realistic = [item for item in scored if item['journal'].get('submission_tier') == 'realistic']
    primary = realistic[0] if realistic else (scored[0] if scored else None)
    backups = []
    for item in scored:
        if not primary or item['journal']['id'] == primary['journal']['id']:
            continue
        backups.append(item)
        if len(backups) == 2:
            break
    shortlist_journal_fit_score = 0
    verified_journal_fit_score = 0
    requirements_checked = False
    if primary:
        shortlist_journal_fit_score = min(2, max(1, int(round(primary['score'] / 3)))) if primary['score'] > 0 else 0
        requirements_checked = bool(primary['requirements']['checked'])
        verified_journal_fit_score = primary['requirements']['score'] if requirements_checked else 0
    return {
        'primary': primary,
        'backups': backups,
        'scored': scored[:8],
        'shortlist_journal_fit_score': shortlist_journal_fit_score,
        'verified_journal_fit_score': verified_journal_fit_score,
        'display_journal_fit_score': verified_journal_fit_score if requirements_checked else shortlist_journal_fit_score,
        'requirements_checked': requirements_checked,
    }


def score_draft_readiness(row, scientific_strength, journal_targets, translation_maturity, tasks=None, evidence_bundle_summary=None):
    score = 0
    if scientific_strength >= 3:
        score += 1
    if evidence_chain_complete(row):
        score += 1
    if normalize(row.get('expected_readouts')) or normalize(row.get('next_test')):
        score += 1
    if journal_targets.get('requirements_checked') and journal_targets.get('verified_journal_fit_score', 0) >= 3:
        score += 1
    if evidence_bundle_summary:
        if (
            int(evidence_bundle_summary.get('reference_count', 0) or 0) >= 20
            and int(evidence_bundle_summary.get('claim_count', 0) or 0) >= 5
            and int(evidence_bundle_summary.get('figure_count', 0) or 0) >= 3
            and int(evidence_bundle_summary.get('section_packet_count', 0) or 0) >= 5
        ):
            score += 1
    if tasks:
        open_critical = [task for task in tasks if task.get('critical') and task.get('status') != 'satisfied']
        if open_critical and all(task.get('status') in {'running', 'satisfied'} and normalize(task.get('output_file')) for task in open_critical):
            score += 1
        contradiction_open = any(
            task.get('task_id') == 'resolve_lead_contradiction' and task.get('status') != 'satisfied'
            for task in tasks
        )
        if contradiction_open:
            score = max(0, score - 1)
    if blocker_severity(row) == 'high':
        score = max(0, score - 1)
    if normalize(translation_maturity) == 'bounded':
        score = max(0, score - 1)
    return max(0, min(4, score))


def has_open_critical_tasks(tasks):
    return any(task.get('critical') and task.get('status') != 'satisfied' for task in tasks or [])


def derive_top_blocker(row, tasks=None, contradictions=None):
    tasks = tasks or []
    contradictions = contradictions or []
    blocked_task = next((task for task in tasks if task.get('critical') and task.get('status') == 'blocked'), None)
    if blocked_task:
        return normalize(blocked_task.get('execution_note')) or normalize(blocked_task.get('rationale')) or normalize(blocked_task.get('label'))

    blocking_issue = next((issue for issue in contradictions if issue.get('blocking')), None)
    if blocking_issue:
        return normalize(blocking_issue.get('statement')) or normalize(blocking_issue.get('recommended_bound'))

    running_task = next((task for task in tasks if task.get('critical') and task.get('status') == 'running'), None)
    if running_task:
        return normalize(running_task.get('execution_note')) or normalize(running_task.get('rationale')) or normalize(running_task.get('label'))

    blockers = split_notes(row.get('blockers')) + split_notes(row.get('contradiction_notes')) + split_notes(row.get('evidence_gaps'))
    return blockers[0] if blockers else ''


def manuscript_gate_state(scientific_strength, journal_targets, draft_readiness, translation_maturity, tasks):
    requirements_checked = bool(journal_targets.get('requirements_checked'))
    verified_journal_fit = int(journal_targets.get('verified_journal_fit_score') or 0)
    if (
        scientific_strength >= 3
        and verified_journal_fit >= 3
        and draft_readiness >= 3
        and requirements_checked
        and normalize(translation_maturity) != 'bounded'
        and not has_open_critical_tasks(tasks)
    ):
        return 'ready to build phase 8 pack'
    if scientific_strength >= 2 or journal_targets.get('display_journal_fit_score', 0) >= 2 or draft_readiness >= 2:
        return 'almost ready'
    return 'not ready'


def ready_for_codex_draft(scientific_strength, journal_targets, draft_readiness, tasks, gate_state):
    open_tasks = [task for task in tasks if task['status'] not in {'satisfied'}]
    return (
        gate_state == 'ready to build phase 8 pack'
        and scientific_strength >= 4
        and int(journal_targets.get('verified_journal_fit_score') or 0) >= 4
        and draft_readiness >= 4
        and not open_tasks
    )


def score_bar_payload(score):
    score = max(0, min(4, int(score)))
    return {'score': score, 'max': 4, 'percent': int(round((score / 4) * 100)) if score else 0}


def task_record(task_id, label, rationale, success_signal, critical=True, satisfied=False):
    status = 'satisfied' if satisfied else 'proposed'
    return {
        'task_id': task_id,
        'label': label,
        'status': status,
        'critical': bool(critical),
        'rationale': rationale,
        'success_signal': success_signal,
        'allowed_states': TASK_STATES,
    }


def build_task_ledger(row, candidate, journal_targets):
    tasks = []
    blockers = split_notes(row.get('blockers')) + split_notes(row.get('contradiction_notes')) + split_notes(row.get('evidence_gaps'))
    translation_maturity = candidate['translation_maturity']
    if blockers:
        lead_blocker = blockers[0]
        tasks.append(task_record(
            'resolve_lead_contradiction',
            'Resolve the lead contradiction or evidence tension',
            f'The current lane still carries a material blocker: {lead_blocker}',
            'A future pass reduces or removes the contradiction/blocker language from the candidate ledger.',
            critical=True,
            satisfied=False,
        ))
    if translation_maturity == 'bounded':
        tasks.append(task_record(
            'strengthen_translational_packet',
            'Strengthen Phase 4 beyond bounded translation',
            'The linked translational packet is still bounded, which caps manuscript confidence.',
            'Linked Phase 4 packet becomes usable/actionable or accumulates stronger compound/trial/readout support.',
            critical=True,
            satisfied=False,
        ))
    if not candidate.get('evidence_chain_complete'):
        tasks.append(task_record(
            'complete_evidence_chain',
            'Complete the cross-phase evidence chain',
            'At least one of the Phase 1-5 links is missing from the ranked candidate.',
            'The candidate has non-empty links across Phases 1, 2, 3, 4, and 5.',
            critical=True,
            satisfied=False,
        ))
    if not normalize(row.get('expected_readouts')) and not normalize(row.get('next_test')):
        tasks.append(task_record(
            'define_figure_spine',
            'Define a cleaner figure and readout spine',
            'The candidate needs a clearer figure path before drafting.',
            'Expected readouts or next-test guidance is populated strongly enough for figure planning.',
            critical=False,
            satisfied=False,
        ))
    primary = journal_targets.get('primary')
    if primary:
        journal_name = primary['journal']['name']
        requirements = primary.get('requirements') or {}
        applicable_guidance = journal_specific_profiles(requirements.get('snapshot') or {}, candidate)
        if not requirements.get('checked'):
            tasks.append(task_record(
                'capture_primary_journal_requirements',
                f'Capture live requirements for {journal_name}',
                'Primary journal choice is still a shortlist call. No verified requirements snapshot exists yet.',
                'A stored requirements snapshot exists for the primary journal and the pack can be checked against it.',
                critical=True,
                satisfied=False,
            ))
        else:
            tasks.append(task_record(
                'tighten_primary_journal_fit',
                f'Check article-type and evidence fit for {journal_name}',
                'Primary journal choice is selected and now needs an explicit pass against the stored requirements snapshot.',
                'Journal fit remains high after validating article type, evidence chain, translation maturity, and scientific strength against the stored requirements snapshot.',
                critical=True,
                satisfied=all([
                    requirements.get('article_type_match'),
                    requirements.get('evidence_chain_match'),
                    requirements.get('translation_match'),
                    requirements.get('scientific_strength_match'),
                ]),
            ))
        if applicable_guidance:
            tasks.append(task_record(
                'prepare_journal_specific_rigor_packet',
                f'Prepare the {journal_name} article-type-specific rigor packet',
                'The primary journal carries article-type-specific Transparency, Rigor and Reproducibility expectations that still need a manuscript-facing checklist and response plan.',
                'A journal-specific rigor checklist exists and the main gaps are mapped into an executable manuscript plan.',
                critical=False,
                satisfied=False,
            ))
    else:
        tasks.append(task_record(
            'select_primary_journal',
            'Select a realistic primary journal',
            'No realistic primary journal could be selected from the current registry and shortlist rules.',
            'The candidate has a realistic primary journal identified from the current registry.',
            critical=True,
            satisfied=False,
        ))
    if not tasks:
        tasks.append(task_record(
            'monitor_candidate',
            'Monitor for new evidence shifts',
            'The candidate is already in relatively strong shape; only watch for meaningful changes.',
            'New evidence changes scientific strength or journal fit materially.',
            critical=False,
            satisfied=True,
        ))
    return tasks


def publication_entry_for_candidate(publication_tracker, candidate_id):
    for entry in publication_tracker.get('entries', []) or []:
        if canonical_candidate_id(entry.get('candidate_id')) == canonical_candidate_id(candidate_id):
            return entry
    return {}


def candidate_priority_key(candidate, publication_status=''):
    return (
        0 if candidate.get('is_active_path') else 1,
        0 if normalize(publication_status) in ACTIVE_APPROVAL_STATUSES else 1,
        -candidate['scientific_strength'],
        -candidate['journal_fit'],
        -candidate['draft_readiness'],
        -candidate['family_preference'],
        -float(candidate['family_score']),
        -float(candidate['confidence_score']),
        -float(candidate['value_score']),
        candidate['title'].lower(),
    )


def build_claim_rows(row, candidate):
    return [
        ('Mechanism lane anchor', summarize_list(row.get('linked_phase1_lane_ids')), 'Phase 1 lane mapping establishes the biological lane the manuscript is about.'),
        ('Causal bridge', summarize_list(row.get('linked_phase2_transition_ids')), 'Phase 2 transition shows the directional bridge that gives the paper its mechanistic spine.'),
        ('Progression burden', summarize_list(row.get('linked_phase3_object_ids')), 'Phase 3 objects define the downstream burden or longitudinal consequence.'),
        ('Translational readout', summarize_list(row.get('linked_phase4_packet_ids')), 'Phase 4 packet ties the paper to targets, readouts, and bounded translational logic.'),
        ('Endotype context', summarize_list(row.get('linked_phase5_endotype_ids')), 'Phase 5 packets keep the manuscript tied to cohort or endotype context instead of generic TBI language.'),
    ]


def ordered_unique(values):
    seen = set()
    ordered = []
    for value in values:
        text = normalize(value)
        if not text or text in seen:
            continue
        seen.add(text)
        ordered.append(text)
    return ordered


def extract_pmids_from_text(value):
    return parse_pmid_list(value)


def repo_relative_path(path):
    path = Path(path)
    try:
        return str(path.relative_to(REPO_ROOT))
    except ValueError:
        return str(path)


def artifact_source_path(state, phase_key):
    report_paths = state.get('report_paths', {})
    if phase_key == 'phase1':
        return repo_relative_path(PROCESS_ENGINE_DOC_PATH)
    mapping = {
        'phase2': report_paths.get('transition'),
        'phase3': report_paths.get('progression'),
        'phase4': report_paths.get('translational'),
        'phase5': report_paths.get('cohort'),
        'phase6': report_paths.get('hypothesis') or report_paths.get('idea_briefs'),
    }
    value = mapping.get(phase_key)
    return repo_relative_path(value) if value else ''


def artifact_id_for_row(phase_key, row):
    if phase_key == 'phase1':
        return normalize(row.get('lane_id') or row.get('display_name'))
    if phase_key == 'phase2':
        return normalize(row.get('transition_id') or row.get('display_name'))
    if phase_key == 'phase3':
        return normalize(row.get('object_id') or row.get('display_name'))
    if phase_key == 'phase4':
        return normalize(row.get('lane_id') or row.get('display_name'))
    if phase_key == 'phase5':
        return normalize(row.get('endotype_id') or row.get('packet_id') or row.get('display_name'))
    return normalize(row.get('candidate_id') or row.get('title') or row.get('display_name'))


def artifact_label_for_row(phase_key, row):
    return normalize(row.get('display_name') or row.get('title') or artifact_id_for_row(phase_key, row))


def support_status_for_row(phase_key, row):
    if phase_key == 'phase1':
        return normalize(row.get('lane_status') or row.get('support_status'))
    if phase_key == 'phase4':
        return normalize(row.get('translation_maturity') or row.get('support_status'))
    if phase_key == 'phase5':
        return normalize(row.get('stratification_maturity') or row.get('support_status'))
    return normalize(row.get('support_status') or row.get('maturity_status') or row.get('hypothesis_status'))


def allowed_language_for_status(phase_key, status):
    status = normalize(status)
    if phase_key == 'phase4' and status == 'bounded':
        return 'Use supported mechanistic language, but keep translational claims bounded and non-clinical.'
    if phase_key == 'phase5' and status == 'bounded':
        return 'Use cohort-language as suggestive context only, not as a locked endotype claim.'
    if status in {'supported', 'usable', 'longitudinally_supported', 'established_in_corpus'}:
        return 'Use affirmative language for the mechanism or pattern, while staying within the current artifact scope.'
    if status in {'provisional', 'bounded', 'seeded', 'emergent_from_tbi_corpus'}:
        return 'Use cautious, hypothesis-oriented wording and explicitly note boundedness.'
    return 'Use cautious wording and avoid definitive causal or clinical claims.'


def phase_entries_for_candidate(row, state, indexes):
    entries = []
    phase_specs = [
        ('phase1', row.get('linked_phase1_lane_ids') or [], indexes.get('process', {})),
        ('phase2', row.get('linked_phase2_transition_ids') or [], indexes.get('transition', {})),
        ('phase3', row.get('linked_phase3_object_ids') or [], indexes.get('progression', {})),
        ('phase4', row.get('linked_phase4_packet_ids') or [], indexes.get('translational', {})),
        ('phase5', row.get('linked_phase5_endotype_ids') or [], indexes.get('cohort', {})),
    ]
    for phase_key, ids, lookup in phase_specs:
        for linked_id in ids:
            phase_row = lookup.get(normalize(linked_id)) or {}
            entries.append({
                'phase_key': phase_key,
                'artifact_id': normalize(linked_id),
                'label': artifact_label_for_row(phase_key, phase_row) or pretty_label(linked_id),
                'row': phase_row,
                'support_status': support_status_for_row(phase_key, phase_row),
                'source_path': artifact_source_path(state, phase_key),
            })
    return entries


def candidate_target_labels(phase_entries):
    labels = []
    for entry in phase_entries:
        phase_key = normalize(entry.get('phase_key'))
        phase_row = entry.get('row') or {}
        if phase_key == 'phase4':
            labels.append(normalize(phase_row.get('primary_target')))
            labels.extend(normalize(item) for item in (phase_row.get('challenger_targets') or []))
            labels.extend(normalize(item) for item in (phase_row.get('biomarker_panel') or []))
        elif phase_key == 'phase5':
            labels.extend(normalize(item) for item in (phase_row.get('linked_primary_targets') or []))
    return [label for label in ordered_unique(labels) if label]


def row_matches_candidate_targets(row, targets):
    if not targets:
        return False
    fields = {
        normalize(row.get('entity_label')).upper(),
        normalize(row.get('entity_id')).upper(),
        normalize(row.get('query_seed')).upper(),
    }
    return any(target.upper() in fields for target in targets if target)


def row_matches_candidate_connector(candidate, row, targets):
    if row_matches_candidate_targets(row, targets):
        return True
    haystack = ' '.join([
        normalize(row.get('canonical_mechanism')),
        normalize(row.get('query_seed')),
        normalize(row.get('title')),
    ])
    return connector_match_score(candidate, haystack) >= 2


def candidate_connector_rows(candidate, phase_entries, connector_context):
    enrichment_rows = connector_context.get('enrichment_rows') or []
    manifest_rows = connector_context.get('manifest_rows') or []
    tenx_template_rows = connector_context.get('tenx_template_rows') or []
    targets = candidate_target_labels(phase_entries)
    matched_enrichment = [
        row for row in enrichment_rows if row_matches_candidate_connector(candidate, row, targets)
    ]
    matched_manifest = [
        row for row in manifest_rows if row_matches_candidate_connector(candidate, row, targets)
    ]
    matched_tenx_template = [
        row
        for row in tenx_template_rows
        if connector_match_score(
            candidate,
            ' '.join([normalize(row.get('canonical_mechanism')), normalize(row.get('query_seed'))]),
        ) >= 2
    ]
    return matched_enrichment, matched_manifest, matched_tenx_template


def summarize_connector_sources(rows):
    counts = {}
    for row in rows:
        source = normalize(row.get('connector_source'))
        if not source:
            continue
        counts[source] = counts.get(source, 0) + 1
    return sorted_counter_rows(counts)


def candidate_connector_summary(candidate, phase_entries, connector_context):
    matched_enrichment, matched_manifest, matched_tenx_template = candidate_connector_rows(
        candidate,
        phase_entries,
        connector_context,
    )
    phase4_entry = next((entry for entry in phase_entries if entry.get('phase_key') == 'phase4'), {})
    phase4_row = phase4_entry.get('row') or {}
    phase5_rows = [entry.get('row') or {} for entry in phase_entries if entry.get('phase_key') == 'phase5']

    source_rows = summarize_connector_sources(matched_enrichment)
    source_labels = [row['label'] for row in source_rows]
    latest_attached_at = max(
        (normalize(row.get('retrieved_at')) for row in matched_enrichment if normalize(row.get('retrieved_at'))),
        default='',
    )
    target_rows = [
        row for row in matched_enrichment if normalize(row.get('evidence_tier')) == 'target_association'
    ]
    trial_rows = [
        row for row in matched_enrichment if normalize(row.get('evidence_tier')) == 'trial_landscape'
    ]
    genomics_rows = [
        row
        for row in matched_enrichment
        if normalize(row.get('evidence_tier')) == 'genomics_expression'
        or normalize(row.get('connector_source')) == 'tenx_genomics'
    ]
    supportive_phase5_row = next(
        (
            row
            for row in phase5_rows
            if normalize(row.get('genomics_support_status')) == 'supportive'
        ),
        {},
    )

    target_sources = []
    if target_rows:
        target_sources.append('Open Targets')
    if normalize(phase4_row.get('primary_target')):
        target_sources.append('repo target-seed pack')
    trial_sources = []
    if trial_rows:
        trial_sources.append('ClinicalTrials.gov')
    if normalize((phase4_row.get('trial_support') or {}).get('status')) in {'supported', 'provisional'}:
        trial_sources.append('translational bridge')
    genomics_sources = []
    if genomics_rows:
        genomics_sources.append('10x Genomics')
    if supportive_phase5_row:
        genomics_sources.append('cohort genomics layer')

    tenx_state = 'not_prepared'
    if genomics_rows:
        tenx_state = 'real_evidence_attached'
    elif matched_tenx_template or any(normalize(row.get('requested_connector')) == 'tenx_genomics' for row in matched_manifest):
        tenx_state = 'template_only'

    data_state = 'no_connector_data_attached'
    if matched_enrichment:
        data_state = 'connector_data_attached'
    elif tenx_state == 'template_only':
        data_state = 'template_only'

    phase4_genomics_status = normalize(phase4_row.get('genomics_support_status'))
    phase4_genomics_detail = normalize(phase4_row.get('genomics_support_detail'))
    supportive_phase5_detail = normalize(
        supportive_phase5_row.get('genomics_signature_detail')
        or supportive_phase5_row.get('genomics_support_detail')
    )
    fallback_phase5_detail = next(
        (
            normalize(row.get('genomics_signature_detail') or row.get('genomics_support_detail'))
            for row in phase5_rows
            if normalize(row.get('genomics_signature_detail') or row.get('genomics_support_detail'))
        ),
        '',
    )
    if genomics_rows:
        genomics_status = 'supported'
        genomics_detail = 'Imported genomics-expression evidence is attached through the connector layer.'
    elif phase4_genomics_status == 'supportive':
        genomics_status = 'supported'
        genomics_detail = phase4_genomics_detail or supportive_phase5_detail or 'Genomics support is attached in the current translational packet.'
    elif supportive_phase5_row:
        genomics_status = 'provisional'
        genomics_detail = supportive_phase5_detail or 'Indirect cohort-level genomics support is present, but no clean 10x packet is attached yet.'
    else:
        genomics_status = phase4_genomics_status or 'not_available'
        genomics_detail = phase4_genomics_detail or fallback_phase5_detail or 'no genomics support attached in this build'

    return {
        'data_state': data_state,
        'matched_row_count': len(matched_enrichment),
        'latest_attached_at': latest_attached_at or connector_context.get('latest_retrieved_at') or '',
        'sources': source_rows,
        'source_labels': source_labels,
        'target_context': {
            'status': normalize(phase4_row.get('target_support')) or ('supported' if target_rows else 'not_available'),
            'source': ' + '.join(ordered_unique(target_sources)) or 'No external target context attached',
            'detail': normalize(phase4_row.get('target_rationale')) or 'No target-support detail attached yet.',
        },
        'trial_context': {
            'status': normalize((phase4_row.get('trial_support') or {}).get('status')) or 'not_available',
            'source': ' + '.join(ordered_unique(trial_sources)) or 'No trial connector attached',
            'detail': normalize((phase4_row.get('trial_support') or {}).get('evidence')) or 'no trial attachment surfaced in this build',
        },
        'genomics_context': {
            'status': genomics_status or 'not_available',
            'source': ' + '.join(ordered_unique(genomics_sources)) or 'No imported genomics evidence',
            'detail': genomics_detail,
        },
        'tenx': {
            'state': tenx_state,
            'imported_row_count': len(genomics_rows),
            'template_row_count': len(matched_tenx_template),
            'template_path': relative_path_or_blank(connector_context.get('tenx_template_path')),
        },
    }


def supporting_pmids_for_phase_entry(entry):
    row = entry.get('row') or {}
    pmids = []
    pmids.extend(parse_pmid_list(row.get('anchor_pmids')))
    pmids.extend(extract_pmids_from_text(row.get('example_signals')))
    if entry.get('phase_key') == 'phase5':
        for paper in row.get('supporting_papers') or []:
            pmids.extend(parse_pmid_list(paper.get('pmid')))
    return ordered_unique(pmids)


def add_reference_support(support_map, ordered_pmids, overrides, pmid, *, phase_key, artifact_id, artifact_label, role, note='', support_status='', metadata_override=None):
    pmid = normalize(pmid)
    if not pmid:
        return
    if pmid not in support_map:
        support_map[pmid] = {
            'phases': [],
            'artifact_refs': [],
            'roles': [],
            'notes': [],
            'support_statuses': [],
        }
        ordered_pmids.append(pmid)
    entry = support_map[pmid]
    if phase_key and phase_key not in entry['phases']:
        entry['phases'].append(phase_key)
    artifact_ref = f'{phase_key}:{artifact_id}' if phase_key and artifact_id else artifact_id
    if artifact_ref and artifact_ref not in entry['artifact_refs']:
        entry['artifact_refs'].append(artifact_ref)
    if role and role not in entry['roles']:
        entry['roles'].append(role)
    if support_status and support_status not in entry['support_statuses']:
        entry['support_statuses'].append(support_status)
    note = normalize(note)
    if note and note not in entry['notes']:
        entry['notes'].append(note)
    if metadata_override:
        overrides.setdefault(pmid, {}).update({key: value for key, value in metadata_override.items() if normalize(value)})


def collect_reference_records(candidate, row, phase_entries, corpus_index):
    support_map = {}
    ordered_pmids = []
    overrides = {}

    for pmid in parse_pmid_list(row.get('anchor_pmids')) + parse_pmid_list(row.get('supporting_pmids')):
        add_reference_support(
            support_map,
            ordered_pmids,
            overrides,
            pmid,
            phase_key='phase6',
            artifact_id=candidate['candidate_id'],
            artifact_label=candidate['title'],
            role='candidate_anchor',
            note='Ranked manuscript candidate support.',
            support_status=normalize(row.get('support_status')),
        )

    for entry in phase_entries:
        phase_key = entry['phase_key']
        artifact_id = entry['artifact_id']
        artifact_label = entry['label']
        phase_row = entry.get('row') or {}
        for pmid in supporting_pmids_for_phase_entry(entry):
            add_reference_support(
                support_map,
                ordered_pmids,
                overrides,
                pmid,
                phase_key=phase_key,
                artifact_id=artifact_id,
                artifact_label=artifact_label,
                role=f'{phase_key}_support',
                note=normalize(phase_row.get('support_reason') or phase_row.get('why_it_matters') or phase_row.get('description')),
                support_status=entry.get('support_status'),
            )
        if phase_key == 'phase5':
            for paper in phase_row.get('supporting_papers') or []:
                pmid = normalize(paper.get('pmid'))
                if not pmid:
                    continue
                add_reference_support(
                    support_map,
                    ordered_pmids,
                    overrides,
                    pmid,
                    phase_key=phase_key,
                    artifact_id=artifact_id,
                    artifact_label=artifact_label,
                    role='phase5_supporting_paper',
                    note='Cohort-supporting paper in linked endotype packet.',
                    support_status=entry.get('support_status'),
                    metadata_override={
                        'title': normalize(paper.get('title')),
                        'source_quality_tier': normalize(paper.get('source_quality_tier')),
                    },
                )

    references = []
    for number, pmid in enumerate(ordered_pmids, start=1):
        record = corpus_reference_record(pmid, corpus_index)
        record.update({key: value for key, value in overrides.get(pmid, {}).items() if normalize(value) and not normalize(record.get(key))})
        record['reference_number'] = number
        record['citation_text'] = citation_text(record)
        record['source_phases'] = support_map[pmid]['phases']
        record['artifact_refs'] = support_map[pmid]['artifact_refs']
        record['support_roles'] = support_map[pmid]['roles']
        record['support_notes'] = support_map[pmid]['notes']
        record['support_statuses'] = support_map[pmid]['support_statuses']
        references.append(record)
    return references


def phase_claim_text(entry, row):
    phase_key = entry['phase_key']
    phase_row = entry.get('row') or {}
    label = entry['label']
    if phase_key == 'phase1':
        return normalize(phase_row.get('description') or f'{label} is the lead mechanism lane for this manuscript path.')
    if phase_key == 'phase2':
        return normalize(phase_row.get('statement_text') or f'{label} is the lead causal bridge in the current evidence chain.')
    if phase_key == 'phase3':
        return normalize(phase_row.get('why_it_matters') or phase_row.get('description') or f'{label} is a downstream progression burden linked to the lead mechanism path.')
    if phase_key == 'phase4':
        primary_target = normalize(phase_row.get('primary_target'))
        readouts = summarize_list(phase_row.get('expected_readouts') or [])
        return normalize(
            phase_row.get('lane_descriptor')
            or f'{label} centers on {primary_target or "the current primary target"} with readouts {readouts}.'
        )
    if phase_key == 'phase5':
        biomarkers = summarize_list(phase_row.get('biomarker_profile') or [])
        return normalize(
            phase_row.get('why_it_matters')
            or f'{label} provides cohort and endotype context with biomarkers {biomarkers}.'
        )
    return normalize(row.get('statement') or row.get('why_now'))


def build_claim_support_matrix(candidate, row, phase_entries):
    claims = []
    candidate_claim_pmids = ordered_unique(
        parse_pmid_list(row.get('anchor_pmids'))
        + parse_pmid_list(row.get('supporting_pmids'))
    )
    claims.append({
        'claim_id': 'phase6_candidate_thesis',
        'phase_key': 'phase6',
        'claim_title': 'Candidate manuscript thesis',
        'claim_text': normalize(row.get('statement') or row.get('why_now') or row.get('decision_rationale')),
        'support_status': normalize(row.get('support_status')),
        'allowed_language': allowed_language_for_status('phase6', normalize(row.get('support_status'))),
        'supporting_pmids': candidate_claim_pmids[:20],
        'artifact_refs': [f"phase6:{candidate['candidate_id']}"],
        'contradiction_notes': split_notes(row.get('blockers')),
        'evidence_gaps': [],
        'source_quality_mix': normalize(row.get('source_quality_mix')),
    })
    for entry in phase_entries:
        phase_row = entry.get('row') or {}
        claims.append({
            'claim_id': f"{entry['phase_key']}_{slugify(entry['artifact_id'])}",
            'phase_key': entry['phase_key'],
            'claim_title': entry['label'],
            'claim_text': phase_claim_text(entry, row),
            'support_status': entry.get('support_status'),
            'allowed_language': allowed_language_for_status(entry['phase_key'], entry.get('support_status')),
            'supporting_pmids': supporting_pmids_for_phase_entry(entry)[:16],
            'artifact_refs': [f"{entry['phase_key']}:{entry['artifact_id']}"],
            'contradiction_notes': split_notes(phase_row.get('contradiction_notes')) + split_notes(phase_row.get('disconfirming_evidence')),
            'evidence_gaps': split_notes(phase_row.get('evidence_gaps')),
            'source_quality_mix': normalize(phase_row.get('source_quality_mix')),
        })
    return claims


def contradiction_severity(statement):
    text = normalize(statement).lower()
    if not text:
        return 'low'
    high_tokens = ['contradict', 'no significant correlation', 'bounded', 'no direct', 'no trial', 'no compound', 'caps manuscript confidence']
    medium_tokens = ['abstract-only', 'seeded', 'provisional', 'still needs', 'weighted cautiously', 'incomplete']
    if any(token in text for token in high_tokens):
        return 'high'
    if any(token in text for token in medium_tokens):
        return 'medium'
    return 'low'


def build_contradiction_register(candidate, row, phase_entries, claims):
    claim_ids_by_phase = {}
    for claim in claims:
        claim_ids_by_phase.setdefault(claim['phase_key'], []).append(claim['claim_id'])
    issues = []
    seen = set()

    def add_issue(phase_key, artifact_id, artifact_label, statement):
        text = normalize(statement)
        if not text or text in seen:
            return
        seen.add(text)
        severity = contradiction_severity(text)
        issues.append({
            'issue_id': f"{phase_key}_{slugify(artifact_id)}_{len(issues) + 1}",
            'phase_key': phase_key,
            'artifact_id': artifact_id,
            'artifact_label': artifact_label,
            'statement': text,
            'severity': severity,
            'blocking': severity == 'high',
            'affected_claim_ids': claim_ids_by_phase.get(phase_key, []),
            'recommended_bound': 'State this as a limitation or bounded uncertainty rather than smoothing it away.',
        })

    for note in split_notes(row.get('blockers')):
        add_issue('phase6', candidate['candidate_id'], candidate['title'], note)
    for entry in phase_entries:
        phase_row = entry.get('row') or {}
        for note in (
            split_notes(phase_row.get('contradiction_notes'))
            + split_notes(phase_row.get('evidence_gaps'))
            + split_notes(phase_row.get('disconfirming_evidence'))
            + split_notes(phase_row.get('lane_notes'))
        ):
            add_issue(entry['phase_key'], entry['artifact_id'], entry['label'], note)
    return issues


def build_source_artifact_manifest(candidate, row, state, phase_entries):
    artifacts = [
        {
            'phase_key': 'phase6',
            'artifact_id': candidate['candidate_id'],
            'label': candidate['title'],
            'source_path': artifact_source_path(state, 'phase6'),
            'support_status': normalize(row.get('support_status')),
            'anchor_pmids': parse_pmid_list(row.get('anchor_pmids'))[:20],
        }
    ]
    for entry in phase_entries:
        phase_row = entry.get('row') or {}
        artifacts.append({
            'phase_key': entry['phase_key'],
            'artifact_id': entry['artifact_id'],
            'label': entry['label'],
            'source_path': entry['source_path'],
            'support_status': entry.get('support_status'),
            'anchor_pmids': supporting_pmids_for_phase_entry(entry)[:16],
            'source_quality_mix': normalize(phase_row.get('source_quality_mix')),
        })
    return artifacts


def build_figure_plan(candidate, row, phase_entries, contradictions):
    progression_labels = [entry['label'] for entry in phase_entries if entry['phase_key'] == 'phase3']
    cohort_labels = [entry['label'] for entry in phase_entries if entry['phase_key'] == 'phase5']
    phase_artifacts = [f"{entry['phase_key']}:{entry['artifact_id']}" for entry in phase_entries]
    figures = [
        {
            'figure_id': 'figure_1',
            'title': 'Phase 1 to Phase 5 evidence chain for the manuscript path',
            'purpose': 'Show how the selected manuscript path moves from lane anchor to causal bridge to progression burden to translational packet to endotype context.',
            'source_artifacts': phase_artifacts,
            'allowed_conclusion': 'This figure supports the structure of the evidence chain, not a claim of clinical proof.',
        },
        {
            'figure_id': 'figure_2',
            'title': 'Biomarker and readout panel for the lead translational packet',
            'purpose': 'Show the candidate biomarker panel, expected readouts, time windows, and sample types.',
            'source_artifacts': [f"phase6:{candidate['candidate_id']}"] + [f"{entry['phase_key']}:{entry['artifact_id']}" for entry in phase_entries if entry['phase_key'] == 'phase4'],
            'allowed_conclusion': 'This figure supports the proposed readout spine and experimental logic, not treatment efficacy.',
        },
        {
            'figure_id': 'figure_3',
            'title': 'Downstream burden and endotype context',
            'purpose': f'Show the progression objects ({", ".join(progression_labels) or "not specified"}) and the linked endotypes ({", ".join(cohort_labels) or "not specified"}).',
            'source_artifacts': [f"{entry['phase_key']}:{entry['artifact_id']}" for entry in phase_entries if entry['phase_key'] in {'phase3', 'phase5'}],
            'allowed_conclusion': 'This figure supports context and stratification framing only.',
        },
        {
            'figure_id': 'figure_4',
            'title': 'Evidence boundaries and contradiction map',
            'purpose': 'Show the main contradictions, evidence gaps, and source-quality boundaries that still limit the paper.',
            'source_artifacts': [f"phase6:{candidate['candidate_id']}"] + [f"{item['phase_key']}:{item['artifact_id']}" for item in phase_entries if item['phase_key'] in {'phase2', 'phase4', 'phase5'}],
            'allowed_conclusion': 'This figure is explicitly a limitation and boundedness figure.',
            'contradiction_count': len(contradictions),
        },
    ]
    return figures


def build_section_packets(candidate, row, state, phase_entries, claims, contradictions, references, source_artifacts, corpus_manifest_path=''):
    lead_refs = ', '.join(f"PMID {item['pmid']}" for item in references[:8]) or 'no lead references yet'
    artifact_paths = ', '.join(ordered_unique(item['source_path'] for item in source_artifacts if item.get('source_path')))
    endotypes = ', '.join(entry['label'] for entry in phase_entries if entry['phase_key'] == 'phase5') or 'no linked endotypes'
    progression = ', '.join(entry['label'] for entry in phase_entries if entry['phase_key'] == 'phase3') or 'no linked progression objects'
    methods_lines = [
        '# Methods Packet',
        '',
        '## Evidence Assembly',
        '',
        f'- Candidate assembled from the live Phase 6 ranked slate using `{candidate["candidate_id"]}`.',
        f'- Phase artifacts used: {artifact_paths or "not recorded"}.',
        f'- Corpus manifest used: `{corpus_manifest_path or "not recorded"}`.',
        '- Manuscript package rules: keep mechanistic claims tied to source artifacts, treat bounded translation as non-clinical, and preserve contradiction language rather than smoothing it away.',
        '- For Journal of Neurotrauma drafting, use `journal_specific_rigor_checklist.md` to build the 500-word Transparency, Rigor and Reproducibility Summary rather than improvising those disclosures later.',
        '',
    ]
    intro_lines = [
        '# Introduction Packet',
        '',
        f'- Lead manuscript thesis: {normalize(row.get("statement") or row.get("why_now"))}',
        f'- Lead rationale: {normalize(row.get("decision_rationale") or row.get("rationale"))}',
        f'- Core anchor references: {lead_refs}.',
        '',
    ]
    results_lines = [
        '# Results Packet',
        '',
        f'- Phase 3 burden to carry in the narrative: {progression}.',
        f'- Endotype context to carry in the narrative: {endotypes}.',
        '',
    ]
    for claim in claims:
        results_lines.append(f"## {claim['claim_title']}")
        results_lines.append('')
        results_lines.append(f"- Claim: {claim['claim_text']}")
        results_lines.append(f"- Allowed wording: {claim['allowed_language']}")
        results_lines.append(f"- Support PMIDs: {', '.join(claim['supporting_pmids']) or 'none recorded'}")
        if claim['contradiction_notes']:
            results_lines.append(f"- Contradictions: {' | '.join(claim['contradiction_notes'])}")
        if claim['evidence_gaps']:
            results_lines.append(f"- Evidence gaps: {' | '.join(claim['evidence_gaps'])}")
        results_lines.append('')
    discussion_lines = [
        '# Discussion Packet',
        '',
        f'- Primary journal target: {candidate["journal_targets"].get("primary", {}).get("journal", {}).get("name", "not selected")}.',
        f'- Primary journal requirements checked: {"yes" if candidate["journal_targets"].get("requirements_checked") else "no"}.',
        '- Discussion should distinguish direct TBI support from cross-disease analog support, and keep bounded translational logic clearly separate from intervention claims.',
        '',
    ]
    limitations_lines = [
        '# Limitations Packet',
        '',
        f'- Total contradiction and evidence-boundary items: {len(contradictions)}.',
        '- Use this packet to keep the manuscript honest about contradiction, bounded translation, abstract-only support, and missing trial/compound attachment.',
        '',
    ]
    for issue in contradictions:
        limitations_lines.append(f"- [{issue['severity']}] {issue['statement']}")
    return {
        '00_introduction_packet.md': '\n'.join(intro_lines) + '\n',
        '01_methods_packet.md': '\n'.join(methods_lines) + '\n',
        '02_results_packet.md': '\n'.join(results_lines) + '\n',
        '03_discussion_packet.md': '\n'.join(discussion_lines) + '\n',
        '04_limitations_packet.md': '\n'.join(limitations_lines) + '\n',
    }


def build_claim_support_markdown(claims, references_by_pmid):
    lines = ['# Claim Support Matrix', '']
    for claim in claims:
        lines.extend([
            f"## {claim['claim_title']}",
            '',
            f"- **Phase:** {claim['phase_key']}",
            f"- **Claim:** {claim['claim_text']}",
            f"- **Support status:** {claim['support_status'] or 'not specified'}",
            f"- **Allowed language:** {claim['allowed_language']}",
            f"- **Artifact refs:** {', '.join(claim['artifact_refs']) or 'none'}",
            f"- **Support PMIDs:** {', '.join(claim['supporting_pmids']) or 'none'}",
        ])
        if claim['supporting_pmids']:
            lines.append('- **Reference details:**')
            for pmid in claim['supporting_pmids'][:8]:
                record = references_by_pmid.get(pmid, {})
                lines.append(f"  - {record.get('citation_text') or f'PMID:{pmid}'}")
        if claim['contradiction_notes']:
            lines.append(f"- **Contradictions:** {' | '.join(claim['contradiction_notes'])}")
        if claim['evidence_gaps']:
            lines.append(f"- **Evidence gaps:** {' | '.join(claim['evidence_gaps'])}")
        lines.append('')
    return '\n'.join(lines)


def build_contradiction_markdown(contradictions):
    lines = ['# Contradiction Register', '']
    if not contradictions:
        lines.append('- No contradiction items recorded.')
        return '\n'.join(lines) + '\n'
    for issue in contradictions:
        lines.extend([
            f"## {issue['artifact_label']}",
            '',
            f"- **Phase:** {issue['phase_key']}",
            f"- **Severity:** {issue['severity']}",
            f"- **Blocking:** {'yes' if issue['blocking'] else 'no'}",
            f"- **Statement:** {issue['statement']}",
            f"- **Affected claims:** {', '.join(issue['affected_claim_ids']) or 'none'}",
            f"- **Recommended bound:** {issue['recommended_bound']}",
            '',
        ])
    return '\n'.join(lines)


def build_reference_markdown(references):
    lines = ['# References Bundle', '']
    for record in references:
        lines.append(f"{record['reference_number']}. {record['citation_text']}")
        lines.append(f"   - Source phases: {', '.join(record.get('source_phases') or []) or 'not specified'}")
        lines.append(f"   - Quality: {record.get('source_quality_tier') or 'unknown'}")
        lines.append(f"   - Artifact refs: {', '.join(record.get('artifact_refs') or []) or 'none'}")
        if record.get('support_notes'):
            lines.append(f"   - Why included: {' | '.join(record['support_notes'][:2])}")
        lines.append('')
    return '\n'.join(lines)


def build_figure_plan_markdown(figures):
    lines = ['# Figure and Table Evidence Plan', '']
    for figure in figures:
        lines.extend([
            f"## {figure['figure_id'].replace('_', ' ').title()}: {figure['title']}",
            '',
            f"- **Purpose:** {figure['purpose']}",
            f"- **Source artifacts:** {', '.join(figure.get('source_artifacts') or []) or 'none'}",
            f"- **Allowed conclusion:** {figure['allowed_conclusion']}",
            '',
        ])
    lines.extend([
        '## Recommended Tables',
        '',
        '- **Table 1:** Claim support matrix by phase, claim, support status, and PMIDs.',
        '- **Table 2:** Journal-fit and requirements snapshot for the primary journal plus backups.',
        '- **Table 3:** Contradiction and evidence-boundary register.',
        '',
    ])
    return '\n'.join(lines)


def build_source_artifact_markdown(source_artifacts):
    lines = ['# Source Artifact Manifest', '']
    for artifact in source_artifacts:
        lines.extend([
            f"## {artifact['label']}",
            '',
            f"- **Phase:** {artifact['phase_key']}",
            f"- **Artifact ID:** {artifact['artifact_id']}",
            f"- **Support status:** {artifact.get('support_status') or 'not specified'}",
            f"- **Source path:** `{artifact.get('source_path') or 'not recorded'}`",
            f"- **Anchor PMIDs:** {', '.join(artifact.get('anchor_pmids') or []) or 'none recorded'}",
        ])
        if normalize(artifact.get('source_quality_mix')):
            lines.append(f"- **Source quality mix:** {artifact['source_quality_mix']}")
        lines.append('')
    return '\n'.join(lines)


def build_data_points_markdown(data_points):
    lines = ['# Data Points Bundle', '']
    for key, value in data_points.items():
        label = pretty_label(key)
        if isinstance(value, list):
            rendered = ', '.join(pretty_label(item) if isinstance(item, str) else normalize(item) for item in value if normalize(item))
            lines.append(f"- **{label}:** {rendered or 'not recorded'}")
        else:
            lines.append(f"- **{label}:** {normalize(value) or 'not recorded'}")
    lines.append('')
    return '\n'.join(lines)


def rigor_check_item(item_id, label, status, evidence, recommendation, critical=False):
    return {
        'item_id': item_id,
        'label': label,
        'status': status,
        'evidence': normalize(evidence),
        'recommendation': normalize(recommendation),
        'critical': bool(critical),
    }


def journal_specific_profiles(snapshot, candidate):
    guidance = snapshot.get('article_type_specific_guidance') or {}
    candidate_modes = {normalize(item) for item in candidate.get('manuscript_modes') or [] if normalize(item)}
    applicable = []
    for guidance_id, spec in guidance.items():
        modes = {normalize(item) for item in spec.get('applies_to_manuscript_modes') or [] if normalize(item)}
        if not modes or candidate_modes.intersection(modes):
            applicable.append((guidance_id, spec))
    return applicable


def summarize_journal_specific_rigor(profiles):
    summary = {
        'profile_count': len(profiles),
        'supported_count': 0,
        'partial_count': 0,
        'missing_count': 0,
        'critical_gap_count': 0,
        'top_gaps': [],
    }
    for profile in profiles:
        for item in profile.get('checklist') or []:
            status = normalize(item.get('status'))
            if status == 'supported':
                summary['supported_count'] += 1
            elif status == 'partial':
                summary['partial_count'] += 1
            else:
                summary['missing_count'] += 1
            if item.get('critical') and status != 'supported':
                summary['critical_gap_count'] += 1
                summary['top_gaps'].append(f"{profile['label']}: {item['label']}")
    summary['top_gaps'] = summary['top_gaps'][:6]
    return summary


def build_journal_specific_rigor_checklist(candidate, row, phase_entries, snapshot):
    if not snapshot or not snapshot.get('verified'):
        return {
            'applicable': False,
            'journal_name': normalize(snapshot.get('journal_name')) or '',
            'profiles': [],
            'summary': {
                'profile_count': 0,
                'supported_count': 0,
                'partial_count': 0,
                'missing_count': 0,
                'critical_gap_count': 0,
                'top_gaps': [],
            },
        }

    applicable = journal_specific_profiles(snapshot, candidate)
    if not applicable:
        return {
            'applicable': False,
            'journal_name': normalize(snapshot.get('journal_name')) or '',
            'profiles': [],
            'summary': {
                'profile_count': 0,
                'supported_count': 0,
                'partial_count': 0,
                'missing_count': 0,
                'critical_gap_count': 0,
                'top_gaps': [],
            },
        }

    phase4_row = next((entry.get('row') or {} for entry in phase_entries if entry['phase_key'] == 'phase4'), {})
    has_expected_readouts = bool(row.get('expected_readouts'))
    has_next_test = bool(normalize(row.get('next_test')))
    has_sample_types = bool(row.get('sample_type'))
    has_time_context = bool(row.get('time_window') or row.get('readout_time_horizon'))
    has_biomarker_panel = bool(row.get('biomarker_panel') or row.get('panel_members'))
    has_phase4 = bool(row.get('linked_phase4_packet_ids'))
    has_phase5 = bool(row.get('linked_phase5_endotype_ids'))
    has_progression = bool(row.get('linked_phase3_object_ids'))
    primary_target = normalize(phase4_row.get('primary_target'))
    best_next_experiment = normalize(phase4_row.get('best_next_experiment'))
    disconfirming = split_notes(phase4_row.get('disconfirming_evidence'))

    profiles = []
    for guidance_id, spec in applicable:
        if guidance_id == 'fluid_biomarker_papers':
            checklist = [
                rigor_check_item(
                    'registration_and_plan',
                    'Study registration and analytic plan disclosure',
                    'missing',
                    'The current manuscript pack does not yet name a registry, preregistration record, or pre-specified analytic plan.',
                    'Prepare the Transparency, Rigor and Reproducibility Summary to state the actual registration status and whether the analytic plan was pre-specified.',
                    critical=True,
                ),
                rigor_check_item(
                    'participant_accounting',
                    'Participant and sample accounting',
                    'partial' if has_phase5 and has_expected_readouts else 'missing',
                    'The pack already names endotype context and clinically relevant readouts, but it does not yet enumerate screened, sampled, analyzed, quality-passed, and outcome-assessed participants.',
                    'Add a CONSORT-style accounting block for screened, sampled, analyzed, QC-passed, and clinically assessed participants.',
                    critical=True,
                ),
                rigor_check_item(
                    'sample_handling_and_blinding',
                    'Biofluid sample handling, batching, and blinding',
                    'partial' if has_sample_types or has_time_context else 'missing',
                    'The pack records sample types and the planned readout lane, but it does not yet capture acquisition timing, handling conditions, batch assignment, freeze-thaw history, or blinded QC roles.',
                    'Capture collection timing, subject state, handling/storage conditions, batch assignment, technical failures, and who was blinded during sample handling and quality control.',
                    critical=True,
                ),
                rigor_check_item(
                    'assay_performance',
                    'Assay validation and analytical performance',
                    'partial' if has_biomarker_panel and has_expected_readouts else 'missing',
                    'The biomarker panel and readout spine are defined, but assay CV, LLOQ/LLOD, batch reproducibility, dilutional linearity, specificity, and spike-recovery are not yet in the manuscript pack.',
                    'Name the assay platform, validation status, replicates, CV, quantitation limits, linearity, specificity, and lot-bridging behavior for the primary biomarker measurements.',
                    critical=True,
                ),
                rigor_check_item(
                    'clinical_outcome_validity',
                    'Clinically relevant assessment validity',
                    'partial' if has_phase5 and (has_expected_readouts or has_progression) else 'missing',
                    'The pack ties the biomarker lane to endotypes and downstream burden, but it does not yet document the validity and reliability of the paired clinical assessments.',
                    'State which inclusion criteria and outcome measures are established standards, emerging standards, or still provisional, and cite the validation basis.',
                    critical=True,
                ),
                rigor_check_item(
                    'statistics_and_replication',
                    'Statistics, multiple comparisons, and replication',
                    'missing',
                    'No manuscript-facing declaration yet covers statistical assumptions, multiple-comparison handling, or the status of internal/external replication.',
                    'Add the planned statistical assumptions, multiple-comparison correction path, outlier/missing-data handling, and the current replication or external-validation status.',
                    critical=True,
                ),
                rigor_check_item(
                    'availability',
                    'Data, code, and sample availability',
                    'missing',
                    'The pack has a source manifest, but it does not yet state repository, DOI, code availability, or future-use permissions for biofluid samples.',
                    'Draft the data-availability, analytic-code, and sample-availability statements expected by the JNT biomarker checklist.',
                    critical=True,
                ),
            ]
        else:
            checklist = [
                rigor_check_item(
                    'registration_and_plan',
                    'Study registration and analytic plan disclosure',
                    'missing',
                    'The current manuscript pack does not yet identify a preregistration record or explicit analytic-plan declaration.',
                    'Prepare the Transparency, Rigor and Reproducibility Summary to state whether the mechanistic study was preregistered and whether the analytic plan was pre-specified.',
                    critical=True,
                ),
                rigor_check_item(
                    'subject_accounting_randomization_blinding',
                    'Subject accounting, randomization, and blinding',
                    'missing',
                    'The pack does not yet describe subject accounting, randomization adequacy, or who was blinded during manipulations and analyses.',
                    'Add subject accounting plus explicit randomization and blinding language, or state clearly why those elements were not used.',
                    critical=True,
                ),
                rigor_check_item(
                    'experimental_manipulation_details',
                    'Experimental manipulation and batching details',
                    'partial' if has_phase4 or has_next_test else 'missing',
                    f"The current pack names the mechanistic next test ({normalize(row.get('next_test')) or 'not specified'}) and translational target ({primary_target or 'not specified'}), but it does not yet capture timing, state, devices, batching, or technical-failure details.",
                    'Capture manipulation timing, subject state, devices or assays used, batching rules, storage conditions, and unexpected events for the lead mechanistic experiment.',
                    critical=True,
                ),
                rigor_check_item(
                    'target_engagement_and_specificity',
                    'Target engagement, specificity, and control logic',
                    'partial' if primary_target or best_next_experiment or disconfirming else 'missing',
                    f"The current packet already points at {primary_target or 'the lead target'} and the next experimental move, but target engagement, specificity controls, cross-validation, and positive/negative controls are not yet explicit.",
                    'Add target-engagement readouts, specificity controls, cross-validation, and negative/positive control language to the mechanistic plan.',
                    critical=True,
                ),
                rigor_check_item(
                    'analysis_and_multiple_comparisons',
                    'Primary-analysis assumptions and multiple-comparison management',
                    'missing',
                    'The manuscript pack does not yet state assumptions, outlier rules, missing-data handling, or how multiple comparisons will be controlled.',
                    'Document statistical assumptions, independence handling, outlier and missing-data rules, and the multiple-comparison strategy for the mechanistic analyses.',
                    critical=True,
                ),
                rigor_check_item(
                    'replication_and_validation',
                    'Replication and external validation status',
                    'missing',
                    'The pack does not yet say whether the mechanistic result has internal replication, external validation, or only a planned replication path.',
                    'State whether replication is complete, ongoing, planned, or absent, and register that plan when possible.',
                    critical=True,
                ),
                rigor_check_item(
                    'availability',
                    'Data, code, and sample availability',
                    'missing',
                    'The current pack does not yet declare where data, code, or residual samples will be shared or requested.',
                    'Draft the data-availability, analytic-code, and sample-availability statements expected by the JNT mechanistic checklist.',
                    critical=True,
                ),
            ]
        profile_summary = {
            'supported_count': sum(1 for item in checklist if item['status'] == 'supported'),
            'partial_count': sum(1 for item in checklist if item['status'] == 'partial'),
            'missing_count': sum(1 for item in checklist if item['status'] == 'missing'),
            'critical_gap_count': sum(1 for item in checklist if item['critical'] and item['status'] != 'supported'),
        }
        profiles.append({
            'profile_id': guidance_id,
            'label': spec.get('label') or pretty_label(guidance_id),
            'definition': normalize(spec.get('definition')),
            'must_address': spec.get('must_address') or [],
            'critical_for_readiness': spec.get('critical_for_readiness') or [],
            'checklist': checklist,
            'status': 'supported' if profile_summary['critical_gap_count'] == 0 else 'needs_work',
            'summary': profile_summary,
        })

    summary = summarize_journal_specific_rigor(profiles)
    return {
        'applicable': True,
        'journal_name': normalize(snapshot.get('journal_name')) or '',
        'profiles': profiles,
        'summary': summary,
    }


def build_journal_specific_rigor_markdown(journal_specific_rigor):
    lines = ['# Journal-Specific Rigor Checklist', '']
    if not journal_specific_rigor.get('applicable'):
        lines.append('- No article-type-specific journal rigor checklist applies to this manuscript candidate yet.')
        lines.append('')
        return '\n'.join(lines)
    summary = journal_specific_rigor.get('summary') or {}
    lines.extend([
        f"- **Journal:** {journal_specific_rigor.get('journal_name') or 'not recorded'}",
        f"- **Applicable profiles:** {len(journal_specific_rigor.get('profiles') or [])}",
        f"- **Critical gaps still open:** {summary.get('critical_gap_count', 0)}",
        '',
    ])
    top_gaps = summary.get('top_gaps') or []
    if top_gaps:
        lines.extend(['## Top Gaps', ''])
        for gap in top_gaps[:6]:
            lines.append(f'- {gap}')
        lines.append('')
    for profile in journal_specific_rigor.get('profiles') or []:
        lines.extend([
            f"## {profile['label']}",
            '',
            f"- **Why it applies:** {profile.get('definition') or 'Profile-specific rigor guidance applies to this manuscript mode.'}",
            f"- **Status:** {profile.get('status')}",
            '',
        ])
        for item in profile.get('checklist') or []:
            critical_tag = ' critical' if item.get('critical') else ''
            lines.append(f"- **{item['label']}** [{item['status']}{critical_tag}]")
            lines.append(f"  - Evidence in current pack: {item.get('evidence') or 'not recorded'}")
            lines.append(f"  - Needed next: {item.get('recommendation') or 'not recorded'}")
        lines.append('')
    return '\n'.join(lines)


def rigor_fill_sections(item_id):
    sections = {
        'registration_and_plan': [
            'Transparency, Rigor and Reproducibility Summary',
            'Methods',
        ],
        'participant_accounting': [
            'Methods',
            'Results',
            'Transparency, Rigor and Reproducibility Summary',
        ],
        'sample_handling_and_blinding': ['Methods', 'Transparency, Rigor and Reproducibility Summary'],
        'assay_performance': ['Methods', 'Supplemental methods/table'],
        'clinical_outcome_validity': ['Methods', 'Discussion'],
        'statistics_and_replication': ['Methods', 'Transparency, Rigor and Reproducibility Summary', 'Discussion'],
        'availability': ['Statements and Declarations', 'Data availability'],
        'subject_accounting_randomization_blinding': ['Methods', 'Transparency, Rigor and Reproducibility Summary'],
        'experimental_manipulation_details': ['Methods', 'Figure legend / supplemental methods'],
        'target_engagement_and_specificity': ['Methods', 'Results', 'Discussion'],
        'analysis_and_multiple_comparisons': ['Methods', 'Transparency, Rigor and Reproducibility Summary'],
        'replication_and_validation': ['Results', 'Discussion', 'Transparency, Rigor and Reproducibility Summary'],
    }
    return sections.get(item_id, ['Transparency, Rigor and Reproducibility Summary'])


def rigor_fill_pack_files(item_id):
    base = [
        'journal_specific_rigor_checklist.md',
        'review_memo.md',
        'pack_manifest.json',
    ]
    specific = {
        'registration_and_plan': ['manuscript_brief.md', 'section_packets/01_methods_packet.md'],
        'participant_accounting': ['section_packets/01_methods_packet.md', 'section_packets/02_results_packet.md', 'data_points.md'],
        'sample_handling_and_blinding': ['data_points.md', 'source_artifact_manifest.md', 'section_packets/01_methods_packet.md'],
        'assay_performance': ['data_points.md', 'references.md', 'section_packets/01_methods_packet.md'],
        'clinical_outcome_validity': ['claim_support_matrix.md', 'references.md', 'section_packets/01_methods_packet.md'],
        'statistics_and_replication': ['claim_support_matrix.md', 'contradiction_register.md', 'section_packets/01_methods_packet.md'],
        'availability': ['source_artifact_manifest.md', 'references.md'],
        'subject_accounting_randomization_blinding': ['section_packets/01_methods_packet.md', 'section_packets/02_results_packet.md'],
        'experimental_manipulation_details': ['data_points.md', 'translational_strengthening_plan.md', 'section_packets/01_methods_packet.md'],
        'target_engagement_and_specificity': ['blocking_resolution_plan.md', 'figure_evidence_plan.md', 'section_packets/02_results_packet.md'],
        'analysis_and_multiple_comparisons': ['claim_support_matrix.md', 'contradiction_register.md', 'section_packets/01_methods_packet.md'],
        'replication_and_validation': ['contradiction_register.md', 'review_memo.md', 'section_packets/03_discussion_packet.md'],
    }
    ordered = []
    for item in base + specific.get(item_id, []):
        if item not in ordered:
            ordered.append(item)
    return ordered


def rigor_fill_questions(profile_id, item_id, candidate, row, phase_entries):
    primary_target = next(
        (
            normalize(entry.get('row', {}).get('primary_target'))
            for entry in phase_entries
            if entry['phase_key'] == 'phase4' and normalize(entry.get('row', {}).get('primary_target'))
        ),
        '',
    )
    next_test = normalize(row.get('next_test') or 'the lead experiment')
    title = normalize(candidate.get('title') or 'the manuscript')
    prompts = {
        'registration_and_plan': [
            f"Was {title} or its underlying study preregistered anywhere, even internally or retrospectively?",
            'Was the analytic plan fixed before looking at the main biomarker outcomes, and can we say that explicitly?',
            'If there was no preregistration, what bounded language will we use so the transparency summary is honest?',
        ],
        'participant_accounting': [
            'How many participants or samples were screened, collected, analyzed, quality-passed, and linked to outcomes?',
            'What exclusion, assay-failure, or missing-data counts need to be stated explicitly?',
            'Can we produce one accounting table that matches the biomarker and clinical readouts exactly?',
        ],
        'sample_handling_and_blinding': [
            'What was the timing of sample collection relative to injury and follow-up state?',
            'How were samples handled, stored, batched, and protected from freeze-thaw variability?',
            'Who was blinded during collection, assay processing, quality control, and analysis?',
        ],
        'assay_performance': [
            'Which assay platform generated the core BBB biomarker readouts, and what validation details are available?',
            'What CV, LLOD/LLOQ, dilutional linearity, specificity, or lot-bridging data can we report?',
            'If formal validation is thin, what exactly can we say without overstating assay certainty?',
        ],
        'clinical_outcome_validity': [
            'Which clinical assessments are treated as established standards versus provisional surrogates?',
            'What citations support the validity of those outcome measures in TBI or adjacent cohorts?',
            'What wording keeps the biomarker-to-outcome bridge bounded but still useful?',
        ],
        'statistics_and_replication': [
            'What were the core statistical assumptions and how were multiple comparisons handled?',
            'How were outliers, missing data, and repeated measures treated?',
            'Do we have internal replication, external validation, or only a planned replication path?',
        ],
        'availability': [
            'What data can be shared now, and what must remain restricted for ethics, privacy, or licensing reasons?',
            'Can we point to code, derived tables, or a future repository plan even if raw data are limited?',
            'What sample-availability statement is honest for this paper?',
        ],
        'subject_accounting_randomization_blinding': [
            'How many subjects or samples entered, completed, failed QC, or dropped out of the mechanistic packet?',
            'Was any randomization used, and if not, what is the clean rationale?',
            'Who was blinded during manipulation, measurement, and analysis?',
        ],
        'experimental_manipulation_details': [
            f"What exact timing, state, batching, and handling details apply to {next_test}?",
            'Which devices, reagents, or assay conditions need to be frozen into Methods so the experiment is reproducible?',
            'What technical failures or unexpected events must be named explicitly?',
        ],
        'target_engagement_and_specificity': [
            f"What evidence would show {primary_target or 'the lead target'} is truly engaged rather than just moving with correlated noise?",
            'Which positive controls, negative controls, or cross-validation readouts should be named in the packet?',
            'How do we show barrier restoration is linked to permeability and downstream spillover rather than marker movement alone?',
        ],
        'analysis_and_multiple_comparisons': [
            'Which analyses are primary versus exploratory for the mechanistic packet?',
            'How will non-independence, multiple comparisons, and missing data be handled explicitly?',
            'What statistical language is safe enough for JNT without pretending the design is stronger than it is?',
        ],
        'replication_and_validation': [
            'Do we already have replication, cross-cohort confirmation, or orthogonal validation for the lead mechanism?',
            'If not, what is the concrete next replication step and how will we state that transparently?',
            'Which current contradiction rows should be reframed as validation debt rather than ignored?',
        ],
    }
    return prompts.get(
        item_id,
        [
            f"What exactly is missing for {profile_id} / {item_id} in {title}?",
            'What can be stated now from the current pack?',
            'What still needs confirmation from source artifacts or author knowledge?',
        ],
    )


def rigor_fill_draft_stub(profile_label, item, candidate, row, phase_entries):
    primary_target = next(
        (
            normalize(entry.get('row', {}).get('primary_target'))
            for entry in phase_entries
            if entry['phase_key'] == 'phase4' and normalize(entry.get('row', {}).get('primary_target'))
        ),
        '',
    )
    next_test = normalize(row.get('next_test') or 'the lead packet')
    item_id = item.get('item_id')
    stubs = {
        'registration_and_plan': 'State whether the study or analysis plan was preregistered, and if it was not, say that explicitly while bounding any post hoc analyses.',
        'participant_accounting': 'Report the counts for screened, sampled, analyzed, quality-passed, and outcome-linked participants in one traceable flow.',
        'sample_handling_and_blinding': 'Name collection timing, handling/storage conditions, batching, and who was blinded during sample processing and analysis.',
        'assay_performance': 'Describe the assay platform, validation state, and the analytical-performance details that support the main biomarker readouts.',
        'clinical_outcome_validity': 'State which paired clinical assessments are established standards versus provisional readouts and cite their validation basis.',
        'statistics_and_replication': 'Define the primary analyses, multiple-comparison handling, outlier/missing-data rules, and the current replication status.',
        'availability': 'Provide a direct data/code/sample availability statement that says what can be shared now and what cannot.',
        'subject_accounting_randomization_blinding': 'Report subject accounting, randomization logic, and who was blinded during the mechanistic experiment.',
        'experimental_manipulation_details': f'Freeze the key reproducibility details for {next_test}, including timing, batching, storage, and technical failures.',
        'target_engagement_and_specificity': f'Explain how {primary_target or "the lead mechanism"} will be shown to be engaged with specificity controls rather than correlated marker movement alone.',
        'analysis_and_multiple_comparisons': 'Spell out the primary analyses, statistical assumptions, and multiple-comparison rules for the mechanistic packet.',
        'replication_and_validation': 'State whether the current mechanistic result is replicated, externally validated, only planned, or still missing validation.',
    }
    return (
        f"{profile_label}: {stubs.get(item_id, item.get('recommendation') or 'Fill this gap explicitly in the manuscript package.')}"
    )


def build_journal_specific_rigor_fill_packet(candidate, row, phase_entries, journal_specific_rigor):
    if not journal_specific_rigor.get('applicable'):
        return {
            'applicable': False,
            'journal_name': journal_specific_rigor.get('journal_name') or '',
            'entries': [],
            'summary': {
                'entry_count': 0,
                'critical_entry_count': 0,
                'open_entry_count': 0,
            },
        }

    entries = []
    for profile in journal_specific_rigor.get('profiles') or []:
        for item in profile.get('checklist') or []:
            if normalize(item.get('status')) == 'supported':
                continue
            entries.append({
                'entry_id': f"{profile['profile_id']}::{item['item_id']}",
                'profile_id': profile['profile_id'],
                'profile_label': profile['label'],
                'gap_label': item['label'],
                'status': item['status'],
                'critical': bool(item.get('critical')),
                'current_pack_evidence': item.get('evidence') or '',
                'needed_next': item.get('recommendation') or '',
                'manuscript_sections': rigor_fill_sections(item.get('item_id')),
                'recommended_pack_files': rigor_fill_pack_files(item.get('item_id')),
                'questions_to_answer': rigor_fill_questions(profile['profile_id'], item.get('item_id'), candidate, row, phase_entries),
                'draft_stub': rigor_fill_draft_stub(profile['label'], item, candidate, row, phase_entries),
            })

    entries.sort(
        key=lambda entry: (
            0 if entry['critical'] else 1,
            0 if normalize(entry.get('status')) == 'missing' else 1,
            entry['profile_label'],
            entry['gap_label'],
        )
    )
    return {
        'applicable': True,
        'journal_name': journal_specific_rigor.get('journal_name') or '',
        'entries': entries,
        'summary': {
            'entry_count': len(entries),
            'critical_entry_count': sum(1 for entry in entries if entry['critical']),
            'open_entry_count': sum(1 for entry in entries if normalize(entry.get('status')) != 'supported'),
        },
    }


def build_journal_specific_rigor_fill_markdown(fill_packet):
    lines = ['# Journal-Specific Rigor Fill Packet', '']
    if not fill_packet.get('applicable'):
        lines.append('- No journal-specific fill packet is needed for this manuscript candidate yet.')
        lines.append('')
        return '\n'.join(lines)
    summary = fill_packet.get('summary') or {}
    lines.extend([
        f"- **Journal:** {fill_packet.get('journal_name') or 'not recorded'}",
        f"- **Open entries to fill:** {summary.get('open_entry_count', 0)}",
        f"- **Critical entries:** {summary.get('critical_entry_count', 0)}",
        '',
        'Use this packet as the working sheet for closing JNT-specific rigor debt. Each entry tells you where the text belongs, what to consult in the pack, and a safe first drafting stub.',
        '',
    ])
    for entry in fill_packet.get('entries') or []:
        critical_tag = 'critical' if entry.get('critical') else 'non-critical'
        lines.extend([
            f"## {entry['gap_label']}",
            '',
            f"- **Profile:** {entry['profile_label']}",
            f"- **Status:** {entry.get('status')}",
            f"- **Priority:** {critical_tag}",
            f"- **Manuscript sections:** {', '.join(entry.get('manuscript_sections') or [])}",
            f"- **Current pack evidence:** {entry.get('current_pack_evidence') or 'not recorded'}",
            f"- **Needed next:** {entry.get('needed_next') or 'not recorded'}",
            '',
            '### Pack files to consult',
            '',
        ])
        for path in entry.get('recommended_pack_files') or []:
            lines.append(f'- `{path}`')
        lines.extend([
            '',
            '### Questions to answer',
            '',
        ])
        for question in entry.get('questions_to_answer') or []:
            lines.append(f'- {question}')
        lines.extend([
            '',
            '### Draft starter',
            '',
            f"> {entry.get('draft_stub') or 'Draft this disclosure explicitly rather than leaving it implicit.'}",
            '',
        ])
    return '\n'.join(lines)


def build_evidence_bundle(candidate, row, state, phase_entries, corpus_index):
    references = collect_reference_records(candidate, row, phase_entries, corpus_index)
    references_by_pmid = {record['pmid']: record for record in references}
    claims = build_claim_support_matrix(candidate, row, phase_entries)
    contradiction_register = build_contradiction_register(candidate, row, phase_entries, claims)
    source_artifacts = build_source_artifact_manifest(candidate, row, state, phase_entries)
    figures = build_figure_plan(candidate, row, phase_entries, contradiction_register)
    section_packets = build_section_packets(
        candidate,
        row,
        state,
        phase_entries,
        claims,
        contradiction_register,
        references,
        source_artifacts,
        corpus_manifest_path=corpus_index.get('manifest_path', ''),
    )

    supporting_claims = {}
    for claim in claims:
        for pmid in claim.get('supporting_pmids') or []:
            supporting_claims.setdefault(pmid, []).append(claim['claim_id'])
    for record in references:
        record['supporting_claim_ids'] = supporting_claims.get(record['pmid'], [])

    journal_snapshot = {
        'primary': candidate['journal_targets'].get('primary'),
        'backups': candidate['journal_targets'].get('backups'),
        'requirements_checked': candidate['journal_targets'].get('requirements_checked'),
        'shortlist_journal_fit_score': candidate['journal_targets'].get('shortlist_journal_fit_score', 0),
        'verified_journal_fit_score': candidate['journal_targets'].get('verified_journal_fit_score', 0),
    }
    primary_snapshot = ((candidate['journal_targets'].get('primary') or {}).get('requirements') or {}).get('snapshot') or {}
    journal_specific_rigor = build_journal_specific_rigor_checklist(candidate, row, phase_entries, primary_snapshot)
    journal_specific_rigor_fill_packet = build_journal_specific_rigor_fill_packet(candidate, row, phase_entries, journal_specific_rigor)
    data_points = {
        'biomarker_panel': row.get('biomarker_panel') or row.get('panel_members') or [],
        'expected_readouts': row.get('expected_readouts') or [],
        'sample_type': row.get('sample_type') or [],
        'time_window': row.get('time_window') or [],
        'readout_time_horizon': row.get('readout_time_horizon') or [],
        'expected_direction': row.get('expected_direction') or [],
        'linked_endotypes': row.get('linked_phase5_endotype_ids') or [],
        'linked_translational_packets': row.get('linked_phase4_packet_ids') or [],
        'primary_target': next((entry.get('row', {}).get('primary_target') for entry in phase_entries if entry['phase_key'] == 'phase4' and normalize(entry.get('row', {}).get('primary_target'))), ''),
        'challenger_targets': next((entry.get('row', {}).get('challenger_targets') for entry in phase_entries if entry['phase_key'] == 'phase4' and entry.get('row', {}).get('challenger_targets')), []),
    }
    summary = {
        'reference_count': len(references),
        'claim_count': len(claims),
        'contradiction_count': len(contradiction_register),
        'figure_count': len(figures),
        'section_packet_count': len(section_packets),
        'primary_journal_verified': bool(candidate['journal_targets'].get('requirements_checked')),
        'corpus_manifest_path': corpus_index.get('manifest_path', ''),
        'journal_specific_rigor_profile_count': journal_specific_rigor.get('summary', {}).get('profile_count', 0),
        'journal_specific_rigor_critical_gaps': journal_specific_rigor.get('summary', {}).get('critical_gap_count', 0),
        'journal_specific_rigor_fill_entries': journal_specific_rigor_fill_packet.get('summary', {}).get('open_entry_count', 0),
    }
    return {
        'references': references,
        'claims': claims,
        'contradictions': contradiction_register,
        'source_artifacts': source_artifacts,
        'figures': figures,
        'section_packets': section_packets,
        'journal_snapshot': journal_snapshot,
        'journal_specific_rigor': journal_specific_rigor,
        'journal_specific_rigor_fill_packet': journal_specific_rigor_fill_packet,
        'data_points': data_points,
        'summary': summary,
        'reference_markdown': build_reference_markdown(references),
        'claim_support_markdown': build_claim_support_markdown(claims, references_by_pmid),
        'contradiction_markdown': build_contradiction_markdown(contradiction_register),
        'figure_plan_markdown': build_figure_plan_markdown(figures),
        'source_artifact_markdown': build_source_artifact_markdown(source_artifacts),
        'data_points_markdown': build_data_points_markdown(data_points),
        'journal_specific_rigor_markdown': build_journal_specific_rigor_markdown(journal_specific_rigor),
        'journal_specific_rigor_fill_markdown': build_journal_specific_rigor_fill_markdown(journal_specific_rigor_fill_packet),
    }


def task_markdown(task, sections):
    lines = [
        f"# {task['label']}",
        '',
        f"- **Task ID:** {task['task_id']}",
        f"- **Status:** {task['status']}",
        f"- **Critical:** {'yes' if task.get('critical') else 'no'}",
        f"- **Rationale:** {task['rationale']}",
        f"- **Success signal:** {task['success_signal']}",
    ]
    if normalize(task.get('execution_note')):
        lines.append(f"- **Executor note:** {task['execution_note']}")
    lines.append('')
    lines.extend(sections)
    lines.append('')
    return '\n'.join(lines)


def execute_manuscript_tasks(candidate, row, phase_entries, evidence_bundle):
    executed = []
    outputs = {}
    primary = candidate.get('journal_targets', {}).get('primary') or {}
    primary_journal = primary.get('journal', {}) or {}
    requirements = primary.get('requirements') or {}
    journal_specific_rigor = evidence_bundle.get('journal_specific_rigor') or {}
    journal_specific_fill_packet = evidence_bundle.get('journal_specific_rigor_fill_packet') or {}
    journal_specific_summary = journal_specific_rigor.get('summary') or {}
    phase4_entries = [entry for entry in phase_entries if entry['phase_key'] == 'phase4']
    missing_phases = [
        phase_key for phase_key, links in [
            ('phase1', row.get('linked_phase1_lane_ids') or []),
            ('phase2', row.get('linked_phase2_transition_ids') or []),
            ('phase3', row.get('linked_phase3_object_ids') or []),
            ('phase4', row.get('linked_phase4_packet_ids') or []),
            ('phase5', row.get('linked_phase5_endotype_ids') or []),
        ] if not links
    ]
    contradictions = evidence_bundle.get('contradictions') or []
    blocking_contradictions = [item for item in contradictions if item.get('blocking')]
    figures = evidence_bundle.get('figures') or []

    for task in candidate.get('task_ledger') or []:
        current = deepcopy(task)
        task_id = current['task_id']
        sections = []
        status = current['status']
        note = ''

        if task_id == 'resolve_lead_contradiction':
            if not contradictions:
                status = 'satisfied'
                note = 'No contradiction rows are currently attached to this manuscript pack.'
                sections.extend([
                    '## Outcome',
                    '',
                    '- No unresolved contradiction items remain in the current bundle.',
                ])
            else:
                status = 'running'
                lead_statement = normalize(contradictions[0].get('statement')) or normalize(current.get('rationale'))
                note = f"Lead contradiction still active: {lead_statement}. {len(blocking_contradictions)} issue(s) remain blocking."
                sections.extend([
                    '## Current contradiction stack',
                    '',
                ])
                for issue in contradictions[:8]:
                    sections.append(f"- [{issue['severity']}] {issue['statement']}")
                if blocking_contradictions:
                    sections.extend([
                        '',
                        '## Blocking issues to resolve before stronger claims',
                        '',
                    ])
                    for issue in blocking_contradictions[:5]:
                        sections.append(f"- {issue['statement']}")
                        if normalize(issue.get('recommended_bound')):
                            sections.append(f"  - Manuscript bound: {issue['recommended_bound']}")
                sections.extend([
                    '',
                    '## BBB manuscript-safe wording for now',
                    '',
                    '- Keep the core claim tied to permeability and immune-spillover biology, not junction-marker movement alone.',
                    '- Treat downstream burden as plausible and mechanistically linked, but not yet fully hardened across every lane.',
                    '- Present contradiction rows explicitly in limitations rather than smoothing them away in Results.',
                    '',
                    '## Best next evidence to lower the contradiction',
                    '',
                ])
                next_test = normalize(row.get('next_test'))
                expected_readouts = summarize_list(row.get('expected_readouts') or [])
                if next_test:
                    sections.append(f"- Run or prioritize the current next-test path: {next_test}")
                if expected_readouts and expected_readouts != 'None yet':
                    sections.append(f"- Keep the manuscript anchored to readouts that directly test leakage/permeability: {expected_readouts}")
                sections.extend([
                    '',
                    '## Next machine emphasis',
                    '',
                    '- Keep bounded wording tied to the contradiction register until new evidence lowers severity or removes the issue.',
                    '- Prefer evidence that tests whether barrier improvement also changes leakage, inflammatory spillover, or downstream burden in the same lane.',
                ])
        elif task_id == 'strengthen_translational_packet':
            if normalize(candidate.get('translation_maturity')) != 'bounded':
                status = 'satisfied'
                note = 'Phase 4 translation is no longer capped at bounded.'
            elif not phase4_entries:
                status = 'blocked'
                note = 'No linked Phase 4 packet is available to strengthen.'
            else:
                status = 'running'
                lead_phase4 = phase4_entries[0].get('row') or {}
                primary_target = normalize(lead_phase4.get('primary_target')) or 'not specified'
                note = f"Translational packet remains bounded around {primary_target}; strengthening brief updated."
                sections.extend([
                    '## Translational packet status',
                    '',
                    f"- Primary target: {primary_target}",
                    f"- Expected readouts: {summarize_list(lead_phase4.get('expected_readouts') or [])}",
                    f"- Best next experiment: {normalize(lead_phase4.get('best_next_experiment')) or 'not specified'}",
                    f"- Why primary now: {normalize(lead_phase4.get('why_primary_now')) or 'not specified'}",
                ])
                disconfirming = split_notes(lead_phase4.get('disconfirming_evidence'))
                if disconfirming:
                    sections.extend([
                        '',
                        '## Disconfirming evidence to keep visible',
                        '',
                    ])
                    sections.extend([f"- {item}" for item in disconfirming[:5]])
                sections.extend([
                    '',
                    '## What would move this packet beyond bounded',
                    '',
                    '- Add direct support that the lead target/readout set shifts the same biology named in the manuscript thesis.',
                    '- Prefer evidence that connects target movement to permeability, inflammatory spillover, or downstream burden rather than isolated marker motion.',
                    '- Keep the manuscript framed as a bounded mechanistic biomarker paper until compound or trial attachment becomes stronger.',
                    '',
                    '## Journal-specific emphasis from Neurotrauma guidance',
                    '',
                    '- Show how the lead experiment will document target engagement, assay reliability, and the validity of the paired mechanistic readouts.',
                    '- Make the bounded translation explicit, but still state what replication, controls, and downstream validation would move the story forward.',
                    '',
                    '## Manuscript-safe translational framing',
                    '',
                    '- Treat the packet as a translational anchor for readout selection, not as proof of intervention-ready efficacy.',
                    '- Use Phase 4 to justify why the panel is worth prioritizing, while keeping actionability language explicitly limited.',
                ])
        elif task_id == 'complete_evidence_chain':
            if not missing_phases:
                status = 'satisfied'
                note = 'All phases 1 through 5 are linked for this manuscript candidate.'
            else:
                status = 'running'
                note = f"Evidence chain still missing: {', '.join(pretty_label(item) for item in missing_phases)}."
                sections.extend([
                    '## Missing phase links',
                    '',
                ])
                sections.extend([f"- {pretty_label(item)}" for item in missing_phases])
        elif task_id == 'define_figure_spine':
            if figures and (normalize(row.get('expected_readouts')) or normalize(row.get('next_test'))):
                status = 'satisfied'
                note = 'Figure spine and readout path are present in the current pack.'
            else:
                status = 'running'
                note = f"Generated figure spine with {len(figures)} planned figure(s), but readout guidance is still sparse."
            sections.extend([
                '## Figure spine',
                '',
            ])
            sections.extend([f"- {figure['title']}" for figure in figures[:6]] or ['- No figure plan was generated.'])
        elif task_id == 'capture_primary_journal_requirements':
            if requirements.get('checked'):
                status = 'satisfied'
                note = f"Verified requirements snapshot exists for {primary_journal.get('name') or 'the primary journal'}."
            else:
                status = 'blocked'
                snapshot = requirements.get('snapshot') or {}
                reason = normalize(snapshot.get('notes')) or 'No verified requirements snapshot is available yet.'
                note = reason
                sections.extend([
                    '## Requirements capture status',
                    '',
                    f"- Journal: {primary_journal.get('name') or 'not selected'}",
                    f"- Snapshot path: `{requirements.get('requirements_path') or 'not recorded'}`",
                    f"- Blocking reason: {reason}",
                ])
        elif task_id == 'tighten_primary_journal_fit':
            if not requirements.get('checked'):
                status = 'blocked'
                note = 'Cannot tighten journal fit until a verified requirements snapshot exists.'
            else:
                matches = {
                    'article_type_match': bool(requirements.get('article_type_match')),
                    'evidence_chain_match': bool(requirements.get('evidence_chain_match')),
                    'translation_match': bool(requirements.get('translation_match')),
                    'scientific_strength_match': bool(requirements.get('scientific_strength_match')),
                }
                status = 'satisfied' if all(matches.values()) else 'running'
                note = f"Requirements check for {primary_journal.get('name') or 'the primary journal'} is {'fully aligned' if all(matches.values()) else 'still mixed'}."
                sections.extend([
                    '## Journal fit checks',
                    '',
                    f"- Article type match: {'yes' if matches['article_type_match'] else 'no'}",
                    f"- Evidence chain match: {'yes' if matches['evidence_chain_match'] else 'no'}",
                    f"- Translation match: {'yes' if matches['translation_match'] else 'no'}",
                    f"- Scientific strength match: {'yes' if matches['scientific_strength_match'] else 'no'}",
                ])
                if requirements.get('reasons'):
                    sections.extend(['', '## Notes', ''])
                    sections.extend([f"- {reason}" for reason in requirements['reasons'][:8]])
                if journal_specific_summary.get('profile_count'):
                    sections.extend([
                        '',
                        '## Article-type-specific rigor watchlist',
                        '',
                        f"- Applicable rigor profiles: {journal_specific_summary.get('profile_count', 0)}",
                        f"- Critical rigor gaps still open: {journal_specific_summary.get('critical_gap_count', 0)}",
                    ])
                    for gap in journal_specific_summary.get('top_gaps', [])[:5]:
                        sections.append(f"- {gap}")
                    sections.append('- Full checklist: `journal_specific_rigor_checklist.md`')
                    sections.append('- Fill packet: `journal_specific_rigor_fill_packet.md`')
        elif task_id == 'prepare_journal_specific_rigor_packet':
            if not journal_specific_rigor.get('applicable'):
                status = 'satisfied'
                note = 'No article-type-specific journal rigor checklist applies to this manuscript candidate.'
                sections.extend([
                    '## Outcome',
                    '',
                    '- No journal-specific rigor packet is needed beyond the base author instructions.',
                ])
            else:
                status = 'running' if journal_specific_summary.get('critical_gap_count', 0) else 'satisfied'
                note = (
                    f"Journal-specific rigor checklist is active with {journal_specific_summary.get('critical_gap_count', 0)} critical gap(s) "
                    f"across {journal_specific_summary.get('profile_count', 0)} profile(s)."
                )
                sections.extend([
                    '## Applicable profiles',
                    '',
                ])
                for profile in journal_specific_rigor.get('profiles') or []:
                    sections.append(f"- {profile['label']}: {profile['summary'].get('critical_gap_count', 0)} critical gap(s)")
                top_gaps = journal_specific_summary.get('top_gaps') or []
                if top_gaps:
                    sections.extend([
                        '',
                        '## Top gaps to close next',
                        '',
                    ])
                    for gap in top_gaps[:6]:
                        sections.append(f"- {gap}")
                sections.extend([
                    '',
                    '## Structured fill packet',
                    '',
                    f"- Open fill entries: {journal_specific_fill_packet.get('summary', {}).get('open_entry_count', 0)}",
                    f"- Critical fill entries: {journal_specific_fill_packet.get('summary', {}).get('critical_entry_count', 0)}",
                ])
                for entry in (journal_specific_fill_packet.get('entries') or [])[:4]:
                    sections.append(
                        f"- {entry['gap_label']} -> {', '.join(entry.get('manuscript_sections') or [])}"
                    )
                sections.extend([
                    '- Use `journal_specific_rigor_fill_packet.md` as the working sheet for filling these disclosures instead of treating the checklist as a dead-end report.',
                    '',
                    '## Drafting rule',
                    '',
                    '- Use `journal_specific_rigor_checklist.md` to write the JNT Transparency, Rigor and Reproducibility Summary instead of improvising it late.',
                    '- Use `journal_specific_rigor_fill_packet.md` to decide which manuscript section to edit and which pack files to consult for each gap.',
                ])
        elif task_id == 'select_primary_journal':
            if primary_journal:
                status = 'satisfied'
                note = f"Primary journal selected: {primary_journal.get('name')}."
                sections.extend([
                    '## Primary journal',
                    '',
                    f"- {primary_journal.get('name')}",
                ])
            else:
                status = 'blocked'
                note = 'No realistic primary journal is currently available for this candidate.'
        elif task_id == 'monitor_candidate':
            status = 'satisfied'
            note = 'Candidate is currently in monitor mode only.'
            sections.extend([
                '## Monitoring state',
                '',
                '- No active manuscript hardening task is needed beyond tracking new evidence shifts.',
            ])

        current['status'] = status
        current['execution_note'] = note
        current['last_executed_at'] = iso_now()
        current['output_file'] = f"task_outputs/{task_id}.md"
        outputs[task_id] = task_markdown(current, sections or ['## Status', '', f"- {note or 'No executor note.'}"])
        executed.append(current)

    status_counts = {state: 0 for state in TASK_STATES}
    for task in executed:
        status_counts[task['status']] = status_counts.get(task['status'], 0) + 1
    summary = {
        'updated_at': iso_now(),
        'total_tasks': len(executed),
        'status_counts': status_counts,
        'open_critical_tasks': sum(1 for task in executed if task.get('critical') and task['status'] != 'satisfied'),
        'blocking_tasks': sum(1 for task in executed if task['status'] == 'blocked'),
        'running_tasks': sum(1 for task in executed if task['status'] == 'running'),
    }
    return executed, outputs, summary


def build_review_memo(candidate, row):
    primary = candidate['journal_targets'].get('primary')
    primary_journal = primary['journal']['name'] if primary else 'No primary journal selected yet'
    requirements = primary.get('requirements') if primary else {}
    blockers = split_notes(row.get('blockers')) + split_notes(row.get('contradiction_notes')) + split_notes(row.get('evidence_gaps'))
    journal_specific_summary = candidate.get('journal_specific_rigor_summary') or {}
    weakest = normalize(candidate.get('top_blocker')) or (blockers[0] if blockers else 'No explicit blocker is recorded, but the packet still needs a cautious read.')
    likely_attack = 'Reviewers are likely to press on bounded Phase 4 support and whether the evidence chain justifies the article type.'
    if any('no significant correlation' in item.lower() for item in blockers):
        likely_attack = 'Reviewers are likely to focus on contradiction rows that report no significant correlation or weak cumulative-damage support.'
    elif journal_specific_summary.get('critical_gap_count'):
        likely_attack = 'Reviewers are likely to press on the journal-specific biomarker and mechanistic rigor disclosures, especially sample handling, assay detail, and replication language.'
    heading = 'Why This Is Worth Developing'
    if candidate['manuscript_gate_state'] == 'ready to build phase 8 pack':
        heading = 'Why This May Be Publishable'
    elif candidate['manuscript_gate_state'] == 'not ready':
        heading = 'Why This Is Still Early'
    journal_note = 'Primary journal requirements have not been captured yet, so journal fit is still provisional.'
    if requirements and requirements.get('checked'):
        journal_note = 'Primary journal requirements have been captured locally and checked against the current pack.'
    rigor_note = ''
    if journal_specific_summary.get('profile_count'):
        rigor_note = f"\n- The journal-specific rigor checklist is active with **{journal_specific_summary.get('critical_gap_count', 0)}** critical gap(s) still open in `journal_specific_rigor_checklist.md`."
    return f"""# Review Memo\n\n## {heading}\n- The candidate has a named paper path with scientific strength currently at **{candidate['scientific_strength']} / 4**.\n- The current primary target journal is **{primary_journal}**.\n- {journal_note}{rigor_note}\n\n## Biggest Weakness\n- {weakest}\n\n## Likely Reviewer Attack\n- {likely_attack}\n\n## Check Before Approval\n- Make sure the title, article type, and main claims stay bounded to what the current artifact chain actually proves.\n- Read `blocking_resolution_plan.md`, `translational_strengthening_plan.md`, `journal_specific_rigor_checklist.md`, and `journal_specific_rigor_fill_packet.md` before treating the pack as close to submission-ready.\n"""


def build_claim_evidence_map(candidate, row):
    lines = [
        '# Claim to Evidence Map',
        '',
        f"- **Candidate:** {candidate['title']}",
        f"- **Decision family:** {pretty_label(candidate['family_id'])}",
        f"- **Article types in play:** {', '.join(candidate['article_type_candidates'])}",
        '',
        '| Claim layer | Engine objects | Why it matters |',
        '| --- | --- | --- |',
    ]
    for label, objects, reason in build_claim_rows(row, candidate):
        lines.append(f'| {label} | {objects} | {reason} |')
    lines.extend([
        '',
        '## Expected Readouts',
        '',
        f"- {summarize_list(row.get('expected_readouts') or [])}",
        '',
        '## Current Boundaries',
        '',
    ])
    blockers = split_notes(row.get('blockers')) + split_notes(row.get('contradiction_notes')) + split_notes(row.get('evidence_gaps'))
    if blockers:
        for item in blockers[:5]:
            lines.append(f'- {item}')
    else:
        lines.append('- No explicit blockers are currently recorded in the ranked row.')
    return '\n'.join(lines) + '\n'


def build_journal_fit_markdown(candidate):
    primary = candidate['journal_targets'].get('primary')
    backups = candidate['journal_targets'].get('backups') or []
    scored = candidate['journal_targets'].get('scored') or []
    journal_specific_summary = candidate.get('journal_specific_rigor_summary') or {}
    lines = ['# Journal Fit', '']
    if primary:
        lines.extend([
            '## Primary Journal',
            '',
            f"- **Name:** {primary['journal']['name']}",
            f"- **Tier:** {pretty_label(primary['journal']['submission_tier'])}",
            f"- **Homepage:** {primary['journal']['homepage']}",
            f"- **Shortlist rationale:** {' '.join(primary['reasons'])}",
            f"- **Requirements checked:** {'Yes' if primary.get('requirements', {}).get('checked') else 'No'}",
            '',
        ])
        if primary.get('requirements', {}).get('reasons'):
            lines.append('### Requirements check')
            lines.append('')
            for reason in primary['requirements']['reasons']:
                lines.append(f'- {reason}')
            requirements_path = normalize(primary['requirements'].get('requirements_path'))
            if requirements_path:
                lines.append(f'- Requirements snapshot: `{requirements_path}`')
            if journal_specific_summary.get('profile_count'):
                lines.append(f"- Article-type-specific rigor profiles: {journal_specific_summary.get('profile_count', 0)}")
                lines.append(f"- Critical rigor gaps still open: {journal_specific_summary.get('critical_gap_count', 0)}")
                if journal_specific_summary.get('top_gaps'):
                    lines.append(f"- Top rigor gaps: {' | '.join(journal_specific_summary.get('top_gaps', [])[:3])}")
                lines.append('- Full checklist: `journal_specific_rigor_checklist.md`')
                lines.append('- Working fill packet: `journal_specific_rigor_fill_packet.md`')
            lines.append('')
    if backups:
        lines.extend(['## Backup Journals', ''])
        for item in backups:
            checked = 'checked' if item.get('requirements', {}).get('checked') else 'shortlist only'
            lines.append(f"- **{item['journal']['name']}** ({pretty_label(item['journal']['submission_tier'])}, {checked}): {' '.join(item['reasons'])}")
        lines.append('')
    lines.extend(['## Scored Journal Pool', ''])
    for item in scored:
        requirement_state = 'checked' if item.get('requirements', {}).get('checked') else 'unchecked'
        lines.append(f"- {item['journal']['name']} — shortlist score {item['score']} ({requirement_state}): {' '.join(item['reasons'])}")
    return '\n'.join(lines) + '\n'


def build_manuscript_brief(candidate, row):
    journal_specific_summary = candidate.get('journal_specific_rigor_summary') or {}
    connector_summary = candidate.get('connector_summary') or {}
    lines = [
        '# Manuscript Brief',
        '',
        f"- **Candidate:** {candidate['title']}",
        f"- **Theme:** {candidate['theme_label']}",
        f"- **Gate state:** {candidate['manuscript_gate_state']}",
        f"- **Scientific strength:** {candidate['scientific_strength']} / 4",
        f"- **Journal fit:** {candidate['journal_fit']} / 4",
        f"- **Draft readiness:** {candidate['draft_readiness']} / 4",
        f"- **Current path:** {'Yes' if candidate['is_active_path'] else 'No'}",
        f"- **Connector data state:** {pretty_label(connector_summary.get('data_state'))}",
        f"- **Connector sources:** {', '.join(connector_summary.get('source_labels') or []) or 'none attached'}",
        '',
        '## Story Spine',
        '',
        f"- **Statement:** {normalize(row.get('statement') or row.get('why_now') or row.get('decision_rationale') or 'No statement emitted.')}",
        f"- **Readouts:** {summarize_list(row.get('expected_readouts') or [])}",
        f"- **Best next test:** {normalize(row.get('next_test') or 'Not yet specified.')}",
        f"- **Journal-specific rigor gaps:** {journal_specific_summary.get('critical_gap_count', 0)} critical gap(s)",
        '',
        '## Improvement Window',
        '',
        '- Default automatic manuscript-improvement passes: 3',
        '- Hard cap: 5 if both scientific strength and journal fit continue to improve',
        '- Manuscript drafting only becomes automatic once the codex handoff trigger is emitted',
        '',
    ]
    return '\n'.join(lines) + '\n'


def build_codex_handoff_markdown(candidate):
    primary = candidate['journal_targets'].get('primary')
    backups = candidate['journal_targets'].get('backups') or []
    journal_specific_summary = candidate.get('journal_specific_rigor_summary') or {}
    lines = [
        '# Codex Draft Handoff',
        '',
        f"- **Candidate ID:** {candidate['candidate_id']}",
        f"- **Candidate title:** {candidate['title']}",
        f"- **Primary journal:** {primary['journal']['name'] if primary else 'Not selected'}",
        f"- **Backups:** {', '.join(item['journal']['name'] for item in backups) if backups else 'None selected yet'}",
        f"- **Article types in play:** {', '.join(candidate['article_type_candidates'])}",
        f"- **Journal-specific rigor gaps:** {journal_specific_summary.get('critical_gap_count', 0)} critical gap(s)",
        '',
        '## Drafting Rules',
        '',
        '- Use the evidence bundle first: references, claim support matrix, contradiction register, source artifact manifest, data points bundle, and section packets.',
        '- Read blocking_resolution_plan.md and translational_strengthening_plan.md before writing Results or Discussion.',
        '- Read journal_specific_rigor_checklist.md before drafting the Transparency, Rigor and Reproducibility Summary or Methods language.',
        '- Use journal_specific_rigor_fill_packet.md to turn each open JNT rigor gap into explicit Methods, Results, Discussion, or Statements language.',
        '- Use the claim-evidence map and review memo before drafting any narrative section.',
        '- Keep all claims bounded to the current artifact chain.',
        '- Re-check the current journal instructions before final formatting.',
        '- Produce the manuscript in markdown first, with journal-aware sectioning.',
        '',
    ]
    return '\n'.join(lines) + '\n'


def build_pack_manifest(candidate, row, publication_entry):
    primary = candidate['journal_targets'].get('primary') or {}
    backups = candidate['journal_targets'].get('backups') or []
    return {
        'updated_at': iso_now(),
        'candidate_id': candidate['candidate_id'],
        'title': candidate['title'],
        'pack_key': candidate['pack_key'],
        'candidate_title': candidate['title'],
        'top_blocker': candidate.get('top_blocker', ''),
        'theme_key': candidate['theme_key'],
        'manuscript_gate_state': candidate['manuscript_gate_state'],
        'scientific_strength': candidate['scientific_strength'],
        'journal_fit': candidate['journal_fit'],
        'journal_shortlist_score': candidate['journal_targets'].get('shortlist_journal_fit_score', 0),
        'journal_verified_fit_score': candidate['journal_targets'].get('verified_journal_fit_score', 0),
        'journal_requirements_checked': candidate['journal_targets'].get('requirements_checked', False),
        'draft_readiness': candidate['draft_readiness'],
        'active_lane': candidate['active_lane'],
        'watchlist_rank': candidate.get('watchlist_rank'),
        'publication_status': normalize(publication_entry.get('status')),
        'primary_journal': primary.get('journal', {}),
        'backup_journals': [item.get('journal', {}) for item in backups],
        'article_type_candidates': candidate['article_type_candidates'],
        'manuscript_modes': candidate['manuscript_modes'],
        'phase_links': {
            'phase1': row.get('linked_phase1_lane_ids') or [],
            'phase2': row.get('linked_phase2_transition_ids') or [],
            'phase3': row.get('linked_phase3_object_ids') or [],
            'phase4': row.get('linked_phase4_packet_ids') or [],
            'phase5': row.get('linked_phase5_endotype_ids') or [],
        },
        'expected_readouts': row.get('expected_readouts') or [],
        'blockers': split_notes(row.get('blockers')),
        'contradiction_notes': split_notes(row.get('contradiction_notes')),
        'evidence_gaps': split_notes(row.get('evidence_gaps')),
        'source_row': row,
        'ready_for_codex_draft': candidate['ready_for_codex_draft'],
        'evidence_bundle': candidate.get('evidence_bundle_summary', {}),
        'connector_summary': candidate.get('connector_summary', {}),
        'journal_specific_rigor': candidate.get('journal_specific_rigor_summary', {}),
        'journal_specific_rigor_fill_packet': {
            'open_entry_count': candidate.get('evidence_bundle_summary', {}).get('journal_specific_rigor_fill_entries', 0),
        },
        'task_execution': candidate.get('task_execution_summary', {}),
        'improvement_window': {
            'default_pass_target': 3,
            'hard_cap': 5,
            'meaningful_improvement_requires': ['scientific_strength', 'journal_fit'],
        },
    }


def recommended_article_type(candidate):
    primary = candidate.get('journal_targets', {}).get('primary') or {}
    overlap = [normalize(item) for item in primary.get('article_type_overlap') or [] if normalize(item)]
    if overlap:
        return overlap[0]
    article_types = [normalize(item) for item in candidate.get('article_type_candidates') or [] if normalize(item)]
    return article_types[0] if article_types else 'review'


def article_type_requirements(snapshot, article_type):
    target = normalize(article_type)
    for key, value in (snapshot.get('article_type_requirements') or {}).items():
        if normalize(key) == target:
            return value
    return {}


def manuscript_draft_title(candidate, row, phase_entries):
    theme_label = normalize(candidate.get('theme_label') or candidate.get('title'))
    family = normalize(candidate.get('family_id'))
    if family == 'most_informative_biomarker_panel':
        return f'{theme_label} as a biomarker-guided manuscript path in traumatic brain injury'
    if family == 'strongest_causal_bridge':
        return f'{theme_label} as a mechanistic bridge in traumatic brain injury'
    if family == 'best_intervention_leverage_point':
        return f'{theme_label} as a bounded translational target in traumatic brain injury'
    if family == 'highest_value_next_task':
        return f'{theme_label} as a methods-guided manuscript path in traumatic brain injury'
    lead_phase4 = next((entry.get('row') or {} for entry in phase_entries if entry.get('phase_key') == 'phase4'), {})
    primary_target = normalize(lead_phase4.get('primary_target'))
    if primary_target:
        return f'{theme_label} with {primary_target}-centered readouts in traumatic brain injury'
    return f'{theme_label} as an evidence-guided manuscript path in traumatic brain injury'


def manuscript_keywords(candidate, row, phase_entries):
    primary_target = next(
        (
            normalize((entry.get('row') or {}).get('primary_target'))
            for entry in phase_entries
            if entry.get('phase_key') == 'phase4' and normalize((entry.get('row') or {}).get('primary_target'))
        ),
        '',
    )
    values = [
        normalize(candidate.get('theme_label')),
        normalize(candidate.get('title')),
        primary_target,
    ]
    values.extend(normalize(item) for item in (row.get('expected_readouts') or [])[:4])
    values.extend(entry.get('label') for entry in phase_entries if entry.get('phase_key') in {'phase3', 'phase5'})
    keywords = []
    for value in values:
        text = normalize(value)
        if not text or text in keywords:
            continue
        keywords.append(text)
        if len(keywords) == 8:
            break
    return keywords


def candidate_draft_folder(root, candidate):
    return Path(root) / candidate['pack_key']


def ready_for_metadata_docx_path(root, candidate):
    return Path(root) / f"{candidate['pack_key']}.docx"


def strip_markdown_formatting(text):
    text = normalize(text)
    if not text:
        return ''
    text = re.sub(r'`([^`]*)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = text.replace('**', '').replace('__', '').replace('*', '').replace('_', '')
    return text.strip()


def apply_inline_markdown_runs(paragraph, text):
    text = normalize(text)
    if not text:
        return
    parts = re.split(r'(\*\*[^*]+\*\*|__[^_]+__|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        bold = (part.startswith('**') and part.endswith('**')) or (part.startswith('__') and part.endswith('__'))
        code = part.startswith('`') and part.endswith('`')
        content = strip_markdown_formatting(part)
        if not content:
            continue
        run = paragraph.add_run(content)
        if bold:
            run.bold = True
        if code:
            run.font.name = 'Courier New'


def write_ready_metadata_docx(manuscript_markdown, output_path, candidate):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError('python-docx is required to build ready-for-metadata DOCX files') from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    if 'Normal' in styles:
        styles['Normal'].font.name = 'Times New Roman'
        styles['Normal'].font.size = Pt(12)
    if 'Title' in styles:
        styles['Title'].font.name = 'Times New Roman'
    if 'Heading 1' in styles:
        styles['Heading 1'].font.name = 'Times New Roman'
    if 'Heading 2' in styles:
        styles['Heading 2'].font.name = 'Times New Roman'
    if 'Heading 3' in styles:
        styles['Heading 3'].font.name = 'Times New Roman'

    lines = manuscript_markdown.splitlines()
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith('# '):
            paragraph = document.add_paragraph(style='Title')
            apply_inline_markdown_runs(paragraph, stripped[2:])
            continue
        if stripped.startswith('## '):
            paragraph = document.add_paragraph(style='Heading 1')
            apply_inline_markdown_runs(paragraph, stripped[3:])
            continue
        if stripped.startswith('### '):
            paragraph = document.add_paragraph(style='Heading 2')
            apply_inline_markdown_runs(paragraph, stripped[4:])
            continue
        if stripped.startswith('#### '):
            paragraph = document.add_paragraph(style='Heading 3')
            apply_inline_markdown_runs(paragraph, stripped[5:])
            continue
        if stripped.startswith('- '):
            paragraph = document.add_paragraph(style='List Bullet')
            apply_inline_markdown_runs(paragraph, stripped[2:])
            continue
        if stripped.startswith('> '):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(strip_markdown_formatting(stripped[2:]))
            run.italic = True
            continue
        paragraph = document.add_paragraph()
        apply_inline_markdown_runs(paragraph, stripped)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def candidate_draft_output_paths(candidate, root=MANUSCRIPT_DRAFTS_ROOT):
    folder = candidate_draft_folder(root, candidate)
    status_path = folder / 'draft_status.md'
    target_journal_path = folder / 'target_journal.md'
    outline_path = folder / 'comprehensive_outline.md'
    missing_items_path = folder / 'missing_metadata_and_items.md'
    manuscript_draft_path = folder / 'manuscript_draft.md'
    manifest_path = folder / GENERATED_DRAFT_MANIFEST_FILENAME
    source_pack_folder = candidate_folder(MANUSCRIPT_ROOT, candidate)
    ready_docx_path = ready_for_metadata_docx_path(READY_FOR_METADATA_DRAFTS_ROOT, candidate)
    return {
        'folder_relative_path': relative_path_or_blank(folder),
        'folder_absolute_path': absolute_path_or_blank(folder),
        'source_pack_relative_path': relative_path_or_blank(source_pack_folder),
        'source_pack_absolute_path': absolute_path_or_blank(source_pack_folder),
        'status_relative_path': relative_path_or_blank(status_path),
        'status_absolute_path': absolute_path_or_blank(status_path),
        'target_journal_relative_path': relative_path_or_blank(target_journal_path),
        'target_journal_absolute_path': absolute_path_or_blank(target_journal_path),
        'outline_relative_path': relative_path_or_blank(outline_path),
        'outline_absolute_path': absolute_path_or_blank(outline_path),
        'missing_items_relative_path': relative_path_or_blank(missing_items_path),
        'missing_items_absolute_path': absolute_path_or_blank(missing_items_path),
        'manuscript_draft_relative_path': relative_path_or_blank(manuscript_draft_path),
        'manuscript_draft_absolute_path': absolute_path_or_blank(manuscript_draft_path),
        'manifest_relative_path': relative_path_or_blank(manifest_path),
        'manifest_absolute_path': absolute_path_or_blank(manifest_path),
        'ready_metadata_docx_relative_path': relative_path_or_blank(ready_docx_path),
        'ready_metadata_docx_absolute_path': absolute_path_or_blank(ready_docx_path),
    }


def operator_metadata_items(candidate):
    primary = candidate.get('journal_targets', {}).get('primary') or {}
    snapshot = (primary.get('requirements') or {}).get('snapshot') or {}
    required_elements = {normalize(item) for item in snapshot.get('required_manuscript_elements') or [] if normalize(item)}
    minimum_keywords = 'Add at least 4 keywords.' if 'minimum 4 keywords' in required_elements else 'Add final keyword set.'
    items = [
        {
            'item_id': 'title_page',
            'label': 'Title page, author list, and affiliations',
            'detail': 'Insert the final title page with author order, affiliations, and corresponding-author contact details.',
        },
        {
            'item_id': 'running_title',
            'label': 'Running title and author metadata',
            'detail': 'Add the short running title, email addresses, ORCID IDs, and any author-note metadata required for submission.',
        },
        {
            'item_id': 'keywords',
            'label': 'Keywords',
            'detail': minimum_keywords,
        },
        {
            'item_id': 'funding',
            'label': 'Funding statement',
            'detail': 'Add the exact funding sources, grant numbers, and sponsor-role statement.',
        },
        {
            'item_id': 'conflicts',
            'label': 'Conflict-of-interest statement',
            'detail': 'Add the final conflict-of-interest declaration for all authors.',
        },
        {
            'item_id': 'author_contributions',
            'label': 'Author contributions',
            'detail': 'Insert the author-contributions section with the final role mapping.',
        },
        {
            'item_id': 'data_availability',
            'label': 'Data and code availability statement',
            'detail': 'Add the final data/code availability statement with repository, DOI, or access process details.',
        },
        {
            'item_id': 'ethics',
            'label': 'Ethics / approvals / consent statement',
            'detail': 'Insert the correct IRB/IACUC, consent, or no-new-human-subjects statement for this manuscript type.',
        },
        {
            'item_id': 'acknowledgements',
            'label': 'Acknowledgements and non-author contributions',
            'detail': 'Add acknowledgements, editorial support disclosures, and any contributor credits.',
        },
        {
            'item_id': 'figure_exports',
            'label': 'Figure files and table exports',
            'detail': 'Export the final figures/tables and align filenames, legends, and callouts with the manuscript draft.',
        },
        {
            'item_id': 'submission_package',
            'label': 'Submission package metadata',
            'detail': 'Prepare the cover letter, suggested reviewers if needed, and journal-form metadata not stored in the engine.',
        },
    ]
    if 'transparency rigor and reproducibility summary' in required_elements:
        items.append({
            'item_id': 'trr_packaging',
            'label': 'Transparency, Rigor and Reproducibility Summary packaging',
            'detail': 'Insert the final journal-positioned Transparency, Rigor and Reproducibility Summary after the conclusion once the content gaps below are closed.',
        })
    return items


def build_missing_metadata_packet(candidate, evidence_bundle):
    metadata_items = operator_metadata_items(candidate)
    open_tasks = []
    for task in candidate.get('task_ledger') or []:
        if normalize(task.get('status')) == 'satisfied' or normalize(task.get('task_id')) == 'monitor_candidate':
            continue
        open_tasks.append({
            'item_id': normalize(task.get('task_id')),
            'label': normalize(task.get('label')),
            'status': normalize(task.get('status')),
            'detail': normalize(task.get('execution_note')) or normalize(task.get('rationale')) or normalize(task.get('success_signal')),
            'critical': bool(task.get('critical')),
        })
    rigor_entries = []
    for entry in (evidence_bundle.get('journal_specific_rigor_fill_packet') or {}).get('entries') or []:
        rigor_entries.append({
            'entry_id': normalize(entry.get('entry_id')),
            'label': normalize(entry.get('gap_label')),
            'status': normalize(entry.get('status')),
            'detail': normalize(entry.get('needed_next')) or normalize(entry.get('current_pack_evidence')),
            'critical': bool(entry.get('critical')),
            'manuscript_sections': entry.get('manuscript_sections') or [],
        })
    content_gap_count = len(open_tasks) + len(rigor_entries)
    ready_for_metadata_only = (
        normalize(candidate.get('manuscript_gate_state')) == 'ready to build phase 8 pack'
        and content_gap_count == 0
    )
    return {
        'metadata_items': metadata_items,
        'content_items': open_tasks,
        'rigor_items': rigor_entries,
        'ready_for_metadata_only': ready_for_metadata_only,
        'summary': {
            'metadata_item_count': len(metadata_items),
            'content_gap_count': content_gap_count,
            'task_gap_count': len(open_tasks),
            'rigor_gap_count': len(rigor_entries),
        },
    }


def build_missing_metadata_markdown(candidate, missing_packet):
    lines = ['# Missing Metadata and Items', '']
    lines.extend([
        f"- **Generated draft status:** {candidate.get('generated_draft_status') or 'draft_generated'}",
        f"- **Ready for metadata only:** {'yes' if missing_packet.get('ready_for_metadata_only') else 'no'}",
        '',
        '## Operator-Supplied Metadata Still Needed',
        '',
    ])
    for item in missing_packet.get('metadata_items') or []:
        lines.append(f"- **{item['label']}:** {item['detail']}")
    lines.append('')
    lines.extend(['## Evidence / Content Gaps Still Open', ''])
    content_items = missing_packet.get('content_items') or []
    if not content_items:
        lines.append('- No open engine-side task gaps remain.')
    else:
        for item in content_items:
            critical_tag = 'critical' if item.get('critical') else 'non-critical'
            lines.append(f"- **{item['label']}** [{item.get('status')}, {critical_tag}]")
            lines.append(f"  - {item.get('detail') or 'No detail recorded.'}")
    lines.append('')
    lines.extend(['## Journal-Specific Rigor Gaps', ''])
    rigor_items = missing_packet.get('rigor_items') or []
    if not rigor_items:
        lines.append('- No journal-specific rigor fill entries remain open.')
    else:
        for item in rigor_items:
            critical_tag = 'critical' if item.get('critical') else 'non-critical'
            lines.append(f"- **{item['label']}** [{item.get('status')}, {critical_tag}]")
            if item.get('manuscript_sections'):
                lines.append(f"  - Manuscript sections: {', '.join(item['manuscript_sections'])}")
            lines.append(f"  - {item.get('detail') or 'No detail recorded.'}")
    lines.append('')
    return '\n'.join(lines)


def generated_draft_status(candidate, missing_packet):
    if missing_packet.get('ready_for_metadata_only'):
        return 'ready_for_metadata_only'
    if missing_packet.get('summary', {}).get('content_gap_count', 0):
        return 'draft_generated_needs_content_or_rigor_fill'
    return 'draft_generated_needs_metadata'


def hydrate_candidate_draft_output(candidate, evidence_bundle, manuscript_drafts_root=MANUSCRIPT_DRAFTS_ROOT):
    missing_packet = build_missing_metadata_packet(candidate, evidence_bundle)
    candidate['missing_metadata_summary'] = missing_packet.get('summary') or {}
    candidate['ready_for_metadata_only'] = bool(missing_packet.get('ready_for_metadata_only'))
    candidate['generated_draft_status'] = generated_draft_status(candidate, missing_packet)
    candidate['draft_output'] = candidate_draft_output_paths(candidate, root=manuscript_drafts_root)
    candidate['draft_output']['generated_draft_status'] = candidate['generated_draft_status']
    candidate['draft_output']['ready_for_metadata_only'] = candidate['ready_for_metadata_only']
    candidate['draft_output']['missing_metadata_item_count'] = candidate['missing_metadata_summary'].get('metadata_item_count', 0)
    candidate['draft_output']['content_gap_count'] = candidate['missing_metadata_summary'].get('content_gap_count', 0)
    return missing_packet


def build_draft_status_markdown(candidate, row, phase_entries, missing_packet):
    primary = candidate.get('journal_targets', {}).get('primary') or {}
    primary_journal = primary.get('journal', {}) or {}
    lines = [
        '# Draft Status',
        '',
        f"- **Title:** {candidate.get('manuscript_draft_title') or candidate.get('title')}",
        f"- **Candidate:** {candidate.get('title')}",
        f"- **Pack key:** {candidate.get('pack_key')}",
        f"- **Generated status:** {candidate.get('generated_draft_status')}",
        f"- **Ready for metadata only:** {'yes' if candidate.get('ready_for_metadata_only') else 'no'}",
        f"- **Gate state:** {candidate.get('manuscript_gate_state')}",
        f"- **Primary journal:** {primary_journal.get('name') or 'not selected'}",
        f"- **Recommended article type:** {candidate.get('recommended_article_type') or 'not specified'}",
        f"- **Scientific strength:** {candidate.get('scientific_strength')} / 4",
        f"- **Journal fit:** {candidate.get('journal_fit')} / 4",
        f"- **Draft readiness:** {candidate.get('draft_readiness')} / 4",
        f"- **Current top blocker:** {candidate.get('top_blocker') or 'none recorded'}",
        f"- **Expected readouts:** {summarize_list(row.get('expected_readouts') or [])}",
        f"- **Next test:** {normalize(row.get('next_test')) or 'not recorded'}",
        '',
        '## Linked Evidence Chain',
        '',
        f"- Phase 1: {summarize_list(row.get('linked_phase1_lane_ids') or [])}",
        f"- Phase 2: {summarize_list(row.get('linked_phase2_transition_ids') or [])}",
        f"- Phase 3: {summarize_list(row.get('linked_phase3_object_ids') or [])}",
        f"- Phase 4: {summarize_list(row.get('linked_phase4_packet_ids') or [])}",
        f"- Phase 5: {summarize_list(row.get('linked_phase5_endotype_ids') or [])}",
        '',
        '## Remaining Work',
        '',
        f"- Operator metadata items: {missing_packet.get('summary', {}).get('metadata_item_count', 0)}",
        f"- Content or rigor gaps: {missing_packet.get('summary', {}).get('content_gap_count', 0)}",
    ]
    if phase_entries:
        lines.extend(['', '## Phase Artifacts In Scope', ''])
        for entry in phase_entries:
            lines.append(f"- {entry['phase_key']}: {entry['label']} ({entry.get('support_status') or 'not specified'})")
    lines.append('')
    return '\n'.join(lines)


def build_target_journal_markdown(candidate):
    primary = candidate.get('journal_targets', {}).get('primary') or {}
    backups = candidate.get('journal_targets', {}).get('backups') or []
    journal = primary.get('journal', {}) or {}
    requirements = primary.get('requirements', {}) or {}
    snapshot = requirements.get('snapshot') or {}
    article_type = candidate.get('recommended_article_type') or recommended_article_type(candidate)
    article_requirements = article_type_requirements(snapshot, article_type)
    lines = ['# Target Journal', '']
    if not journal:
        lines.append('- No primary journal is selected yet.')
        lines.append('')
        return '\n'.join(lines)
    lines.extend([
        f"- **Primary journal:** {journal.get('name')}",
        f"- **Homepage:** {journal.get('homepage') or 'not recorded'}",
        f"- **Submission tier:** {pretty_label(journal.get('submission_tier'))}",
        f"- **Recommended article type:** {article_type or 'not specified'}",
        f"- **Requirements snapshot verified:** {'yes' if requirements.get('checked') else 'no'}",
        '',
        '## Why This Journal',
        '',
    ])
    for reason in primary.get('reasons') or []:
        lines.append(f'- {reason}')
    lines.append('')
    if requirements.get('reasons'):
        lines.extend(['## Requirements Check', ''])
        for reason in requirements.get('reasons') or []:
            lines.append(f'- {reason}')
        requirement_path = normalize(requirements.get('requirements_path'))
        if requirement_path:
            lines.append(f'- Requirements snapshot path: `{requirement_path}`')
        lines.append('')
    if article_requirements:
        lines.extend(['## Article-Type Constraints', ''])
        for key, value in article_requirements.items():
            lines.append(f"- **{pretty_label(key)}:** {value}")
        lines.append('')
    required_elements = snapshot.get('required_manuscript_elements') or []
    if required_elements:
        lines.extend(['## Required Manuscript Elements', ''])
        for item in required_elements:
            lines.append(f'- {item}')
        lines.append('')
    if backups:
        lines.extend(['## Backup Journals', ''])
        for item in backups:
            backup = item.get('journal', {}) or {}
            checked = 'checked' if (item.get('requirements') or {}).get('checked') else 'shortlist only'
            lines.append(f"- **{backup.get('name') or 'Unnamed backup'}** ({checked}): {' '.join(item.get('reasons') or [])}")
        lines.append('')
    return '\n'.join(lines)


def claim_support_suffix(claim):
    pmids = [normalize(item) for item in claim.get('supporting_pmids') or [] if normalize(item)]
    if not pmids:
        return ''
    return f" (PMIDs: {', '.join(pmids[:4])})"


def build_comprehensive_outline_markdown(candidate, row, phase_entries, evidence_bundle, missing_packet):
    figures = evidence_bundle.get('figures') or []
    claims = evidence_bundle.get('claims') or []
    primary = candidate.get('journal_targets', {}).get('primary') or {}
    lines = [
        '# Comprehensive Outline',
        '',
        '## Positioning',
        '',
        f"- **Draft title:** {candidate.get('manuscript_draft_title') or candidate.get('title')}",
        f"- **Primary journal:** {(primary.get('journal') or {}).get('name') or 'not selected'}",
        f"- **Recommended article type:** {candidate.get('recommended_article_type') or 'not specified'}",
        f"- **Core thesis:** {normalize(row.get('statement') or row.get('why_now') or row.get('decision_rationale')) or 'No thesis recorded.'}",
        '',
        '## Front Matter',
        '',
        '- Title page and author metadata',
        '- Unstructured abstract',
        '- Keywords',
        '- Transparency, Rigor and Reproducibility Summary',
        '- Statements and declarations',
        '',
        '## Main Body Outline',
        '',
        '### 1. Introduction',
        '',
        f"- Open with the current thesis: {normalize(row.get('statement') or row.get('why_now') or row.get('decision_rationale')) or 'No thesis recorded.'}",
        f"- Establish the lead mechanism lane(s): {summarize_list(row.get('linked_phase1_lane_ids') or [])}",
        f"- Set up the directional bridge(s): {summarize_list(row.get('linked_phase2_transition_ids') or [])}",
        '',
        '### 2. Results / Evidence Synthesis',
        '',
    ]
    for claim in claims:
        lines.extend([
            f"#### {claim['claim_title']}",
            '',
            f"- Claim to carry: {claim['claim_text']}",
            f"- Allowed wording: {claim['allowed_language']}",
            f"- Support anchors: {', '.join(claim.get('supporting_pmids') or []) or 'none recorded'}",
            '',
        ])
    lines.extend([
        '### 3. Discussion',
        '',
        f"- Frame the journal fit: {(primary.get('journal') or {}).get('name') or 'not selected'}",
        f"- Carry the translational maturity boundary: {candidate.get('translation_maturity') or 'not specified'}",
        f"- Keep the top blocker explicit: {candidate.get('top_blocker') or 'none recorded'}",
        '',
        '### 4. Limitations',
        '',
    ])
    contradictions = evidence_bundle.get('contradictions') or []
    if contradictions:
        for issue in contradictions[:8]:
            lines.append(f"- {issue['statement']}")
    else:
        lines.append('- No contradiction register items are currently attached.')
    lines.extend([
        '',
        '### 5. Statements and Declarations',
        '',
        '- Funding',
        '- Conflict of interest',
        '- Author contributions',
        '- Data availability',
        '- Ethics / approvals',
        '',
        '## Figure and Table Plan',
        '',
    ])
    for figure in figures:
        lines.append(f"- **{figure['title']}:** {figure['purpose']}")
    lines.extend([
        '',
        '## Remaining Work Before Submission Packaging',
        '',
        f"- Metadata items still needed: {missing_packet.get('summary', {}).get('metadata_item_count', 0)}",
        f"- Content or rigor gaps still open: {missing_packet.get('summary', {}).get('content_gap_count', 0)}",
        '',
    ])
    return '\n'.join(lines)


def build_full_manuscript_draft_markdown(candidate, row, phase_entries, evidence_bundle, missing_packet):
    primary = candidate.get('journal_targets', {}).get('primary') or {}
    primary_journal = primary.get('journal', {}) or {}
    claims = evidence_bundle.get('claims') or []
    contradictions = evidence_bundle.get('contradictions') or []
    figures = evidence_bundle.get('figures') or []
    references = evidence_bundle.get('references') or []
    keywords = candidate.get('manuscript_keywords') or []
    statement = normalize(row.get('statement') or row.get('why_now') or row.get('decision_rationale')) or 'No thesis statement is recorded in the current pack.'
    rationale = normalize(row.get('decision_rationale') or row.get('rationale')) or 'No explicit decision rationale is recorded.'
    readouts = summarize_list(row.get('expected_readouts') or [])
    next_test = normalize(row.get('next_test')) or 'not recorded'
    lead_phase1 = next((entry for entry in phase_entries if entry.get('phase_key') == 'phase1'), {})
    lead_phase2 = next((entry for entry in phase_entries if entry.get('phase_key') == 'phase2'), {})
    lead_phase4 = next((entry for entry in phase_entries if entry.get('phase_key') == 'phase4'), {})
    phase3_labels = ', '.join(entry['label'] for entry in phase_entries if entry.get('phase_key') == 'phase3') or 'no linked downstream objects'
    phase5_labels = ', '.join(entry['label'] for entry in phase_entries if entry.get('phase_key') == 'phase5') or 'no linked endotypes'
    primary_target = normalize((lead_phase4.get('row') or {}).get('primary_target')) or 'the linked Phase 4 packet'
    abstract_lines = [
        f"{candidate.get('theme_label') or candidate.get('title')} is the current manuscript path selected from the Phase 8 evidence pack.",
        f"The current thesis is: {statement}",
        f"The evidence chain links {lead_phase1.get('label') or 'the linked Phase 1 lane'} to {lead_phase2.get('label') or 'the linked Phase 2 bridge'}, with downstream burden in {phase3_labels}.",
        f"The translational readout spine centers on {primary_target} with expected readouts {readouts}.",
        f"The manuscript remains intentionally bounded: translation maturity is {candidate.get('translation_maturity') or 'not specified'}, and the leading boundary is {candidate.get('top_blocker') or 'not specified'}.",
    ]
    lines = [
        f"# {candidate.get('manuscript_draft_title') or candidate.get('title')}",
        '',
        '> Draft-first manuscript generated directly from the current Phase 8 evidence pack. Replace bracketed metadata placeholders manually and keep all claims bounded to the cited pack content.',
        '',
        '## Title Page',
        '',
        '- Full title: [Confirm final journal-facing title]',
        '- Running title: [Needed]',
        '- Authors: [Needed]',
        '- Affiliations: [Needed]',
        '- Corresponding author: [Needed]',
        '',
        '## Abstract',
        '',
        ' '.join(abstract_lines),
        '',
        '## Keywords',
        '',
        f"- {', '.join(keywords) if keywords else '[Add final keyword list]'}",
        '',
        '## Introduction',
        '',
        f"{statement} This draft is anchored to {lead_phase1.get('label') or 'the linked Phase 1 lane'} because that lane provides the current biological entry point into the manuscript path.",
        '',
        f"The linked rationale in the Phase 8 pack is straightforward: {rationale} The directional bridge into {lead_phase2.get('label') or 'the linked Phase 2 transition'} gives the manuscript a mechanistic spine, while the linked Phase 3 and Phase 5 artifacts keep the story tied to downstream burden and cohort context rather than to generic TBI language.",
        '',
        '## Results / Evidence Synthesis',
        '',
    ]
    for claim in claims:
        lines.extend([
            f"### {claim['claim_title']}",
            '',
            f"{claim['claim_text']}{claim_support_suffix(claim)}",
            '',
            f"Drafting rule: {claim['allowed_language']}",
            '',
        ])
        if claim.get('contradiction_notes'):
            lines.append(f"Current contradiction notes: {' | '.join(claim['contradiction_notes'])}")
            lines.append('')
        if claim.get('evidence_gaps'):
            lines.append(f"Current evidence gaps: {' | '.join(claim['evidence_gaps'])}")
            lines.append('')
    lines.extend([
        '## Discussion',
        '',
        f"This manuscript is currently targeted to {primary_journal.get('name') or 'the current primary journal'} as a {candidate.get('recommended_article_type') or 'draft manuscript'} because the current evidence pack already carries a named thesis, a linked Phase 1 to Phase 5 chain, and a draft readout spine built around {readouts}.",
        '',
        f"The draft should still remain bounded. Translation maturity is {candidate.get('translation_maturity') or 'not specified'}, the next experimental move is {next_test}, and the current top blocker is {candidate.get('top_blocker') or 'not specified'}. These constraints should remain explicit rather than being smoothed away during prose cleanup.",
        '',
        f"The cohort and endotype context currently in scope is {phase5_labels}. Use that context to keep the narrative specific about where the manuscript path appears strongest and where it remains only provisional.",
        '',
        '## Limitations',
        '',
    ])
    if contradictions:
        for issue in contradictions[:10]:
            lines.append(f"- {issue['statement']}")
    else:
        lines.append('- No explicit contradiction register items are attached in the current pack.')
    lines.extend([
        '',
        '## Transparency, Rigor and Reproducibility Summary Starter',
        '',
    ])
    rigor_entries = (evidence_bundle.get('journal_specific_rigor_fill_packet') or {}).get('entries') or []
    if rigor_entries:
        for entry in rigor_entries[:6]:
            lines.append(f"- {entry.get('draft_stub') or entry.get('needed_next') or entry.get('gap_label')}")
    else:
        lines.append('- No open journal-specific rigor fill packet entries are currently attached to this pack.')
    lines.extend([
        '',
        '## Statements and Declarations',
        '',
        '### Funding',
        '',
        '[Needed]',
        '',
        '### Conflict of Interest',
        '',
        '[Needed]',
        '',
        '### Author Contributions',
        '',
        '[Needed]',
        '',
        '### Data Availability',
        '',
        '[Needed]',
        '',
        '### Ethics / Approvals',
        '',
        '[Needed]',
        '',
        '## Figure and Table Plan',
        '',
    ])
    for figure in figures:
        lines.append(f"- **{figure['title']}:** {figure['purpose']}")
    lines.extend([
        '',
        '## Missing Metadata and Packaging Items',
        '',
        f"- Operator metadata items still needed: {missing_packet.get('summary', {}).get('metadata_item_count', 0)}",
        f"- Content or rigor gaps still open: {missing_packet.get('summary', {}).get('content_gap_count', 0)}",
        f"- See `{Path(candidate.get('draft_output', {}).get('missing_items_relative_path') or '').name or 'missing_metadata_and_items.md'}` for the detailed working list.",
        '',
        '## References',
        '',
    ])
    if references:
        for record in references:
            lines.append(f"{record['reference_number']}. {record['citation_text']}")
    else:
        lines.append('[No references were emitted in the current evidence bundle.]')
    lines.append('')
    return '\n'.join(lines)


def draft_queue_entry(candidate):
    draft_output = candidate.get('draft_output') or {}
    return {
        'candidate_id': candidate.get('candidate_id'),
        'pack_key': candidate.get('pack_key'),
        'title': candidate.get('title'),
        'manuscript_gate_state': candidate.get('manuscript_gate_state'),
        'generated_draft_status': candidate.get('generated_draft_status'),
        'ready_for_metadata_only': bool(candidate.get('ready_for_metadata_only')),
        'active_lane': bool(candidate.get('active_lane')),
        'watchlist_rank': candidate.get('watchlist_rank'),
        'folder_relative_path': draft_output.get('folder_relative_path', ''),
        'folder_absolute_path': draft_output.get('folder_absolute_path', ''),
        'status_relative_path': draft_output.get('status_relative_path', ''),
        'status_absolute_path': draft_output.get('status_absolute_path', ''),
        'target_journal_relative_path': draft_output.get('target_journal_relative_path', ''),
        'target_journal_absolute_path': draft_output.get('target_journal_absolute_path', ''),
        'outline_relative_path': draft_output.get('outline_relative_path', ''),
        'outline_absolute_path': draft_output.get('outline_absolute_path', ''),
        'missing_items_relative_path': draft_output.get('missing_items_relative_path', ''),
        'missing_items_absolute_path': draft_output.get('missing_items_absolute_path', ''),
        'manuscript_draft_relative_path': draft_output.get('manuscript_draft_relative_path', ''),
        'manuscript_draft_absolute_path': draft_output.get('manuscript_draft_absolute_path', ''),
        'ready_metadata_docx_relative_path': draft_output.get('ready_metadata_docx_relative_path', ''),
        'ready_metadata_docx_absolute_path': draft_output.get('ready_metadata_docx_absolute_path', ''),
        'manifest_relative_path': draft_output.get('manifest_relative_path', ''),
        'manifest_absolute_path': draft_output.get('manifest_absolute_path', ''),
        'missing_metadata_item_count': draft_output.get('missing_metadata_item_count', 0),
        'content_gap_count': draft_output.get('content_gap_count', 0),
    }


def build_generated_draft_summary(candidates, root=MANUSCRIPT_DRAFTS_ROOT, ready_root=READY_FOR_METADATA_DRAFTS_ROOT):
    root = Path(root)
    ready_root = Path(ready_root)
    entries = [draft_queue_entry(candidate) for candidate in candidates]
    ready = [entry for entry in entries if entry.get('ready_for_metadata_only')]
    return {
        'root_relative_path': relative_path_or_blank(root),
        'root_absolute_path': absolute_path_or_blank(root),
        'index_relative_path': relative_path_or_blank(root / GENERATED_DRAFT_INDEX_FILENAME),
        'index_absolute_path': absolute_path_or_blank(root / GENERATED_DRAFT_INDEX_FILENAME),
        'ready_metadata_root_relative_path': relative_path_or_blank(ready_root),
        'ready_metadata_root_absolute_path': absolute_path_or_blank(ready_root),
        'entries': entries,
        'ready_for_metadata_only': ready,
    }


def build_candidate_payload(row, state, direction_registry, journal_registry, publication_tracker, indexes, corpus_index, connector_context=None):
    connector_context = connector_context or build_connector_context()
    candidate_id = canonical_candidate_id(row.get('candidate_id') or row.get('title'))
    title = normalize(row.get('title') or row.get('display_name') or candidate_id)
    theme = theme_key(row)
    active_path_id = canonical_candidate_id(direction_registry.get('active_path_id'))
    candidate = {
        'candidate_id': candidate_id,
        'pack_key': pack_key(candidate_id),
        'title': title,
        'theme_key': theme,
        'theme_label': normalize(row.get('display_name') or title),
        'family_id': normalize(row.get('family_id')),
        'family_preference': family_preference(row.get('family_id')),
        'family_score': float(row.get('family_score', 0) or 0),
        'confidence_score': float(row.get('confidence_score', 0) or 0),
        'value_score': float(row.get('value_score', 0) or 0),
        'support_status': normalize(row.get('support_status') or row.get('strength_tag') or 'bounded'),
        'translation_maturity': translation_maturity_for_row(row, indexes),
        'article_type_candidates': candidate_article_types(row),
        'manuscript_modes': candidate_modes(row),
        'audiences': candidate_audiences(row),
        'evidence_chain_complete': evidence_chain_complete(row),
        'is_active_path': active_path_id == candidate_id,
        'last_pack_refresh': state.get('report_timestamps', {}).get('hypothesis') or direction_registry.get('last_updated') or '',
    }
    candidate['scientific_strength'] = score_scientific_strength(row, indexes)
    journal_targets = choose_journal_targets(candidate, journal_registry)
    candidate['journal_targets'] = journal_targets
    candidate['journal_fit'] = journal_targets['display_journal_fit_score']
    phase_entries = phase_entries_for_candidate(row, state, indexes)
    candidate['connector_summary'] = candidate_connector_summary(candidate, phase_entries, connector_context)
    evidence_bundle = build_evidence_bundle(candidate, row, state, phase_entries, corpus_index)
    candidate['evidence_bundle_summary'] = evidence_bundle['summary']
    candidate['journal_specific_rigor_summary'] = evidence_bundle.get('journal_specific_rigor', {}).get('summary', {})
    candidate['task_ledger'] = build_task_ledger(row, candidate, journal_targets)
    candidate['task_ledger'], _, candidate['task_execution_summary'] = execute_manuscript_tasks(candidate, row, phase_entries, evidence_bundle)
    candidate['top_blocker'] = derive_top_blocker(row, candidate['task_ledger'], evidence_bundle.get('contradictions'))
    candidate['draft_readiness'] = score_draft_readiness(
        row,
        candidate['scientific_strength'],
        journal_targets,
        candidate['translation_maturity'],
        tasks=candidate['task_ledger'],
        evidence_bundle_summary=evidence_bundle.get('summary'),
    )
    candidate['scientific_strength_bar'] = score_bar_payload(candidate['scientific_strength'])
    candidate['journal_fit_bar'] = score_bar_payload(candidate['journal_fit'])
    candidate['draft_readiness_bar'] = score_bar_payload(candidate['draft_readiness'])
    candidate['manuscript_gate_state'] = manuscript_gate_state(candidate['scientific_strength'], journal_targets, candidate['draft_readiness'], candidate['translation_maturity'], candidate['task_ledger'])
    publication_entry = publication_entry_for_candidate(publication_tracker, candidate_id)
    candidate['publication_status'] = normalize(publication_entry.get('status'))
    candidate['ready_for_codex_draft'] = ready_for_codex_draft(candidate['scientific_strength'], journal_targets, candidate['draft_readiness'], candidate['task_ledger'], candidate['manuscript_gate_state'])
    candidate['draft_status'] = 'ready_for_codex_draft' if candidate['ready_for_codex_draft'] else ('pack_ready' if candidate['manuscript_gate_state'] == 'ready to build phase 8 pack' else 'gathering evidence')
    candidate['review_memo_status'] = 'available' if candidate['manuscript_gate_state'] != 'not ready' else 'early warning'
    candidate['queue_status'] = 'submitted_track' if candidate['publication_status'] in SUBMISSION_TRACKER_EXIT_STATUSES else ('active_review' if candidate['publication_status'] in ACTIVE_APPROVAL_STATUSES else 'candidate')
    candidate['recommended_article_type'] = recommended_article_type(candidate)
    candidate['manuscript_draft_title'] = manuscript_draft_title(candidate, row, phase_entries)
    candidate['manuscript_keywords'] = manuscript_keywords(candidate, row, phase_entries)
    hydrate_candidate_draft_output(candidate, evidence_bundle)
    return candidate, publication_entry


def representative_candidate_rows(state, direction_registry):
    rows = list(state.get('portfolio') or state.get('hypothesis', {}).get('rows') or [])
    if not rows:
        return []
    groups = {}
    for row in rows:
        groups.setdefault(canonical_candidate_id(row.get('candidate_id') or row.get('title')), []).append(row)
    active_path_id = canonical_candidate_id(direction_registry.get('active_path_id'))

    def row_key(row):
        candidate_id = canonical_candidate_id(row.get('candidate_id') or row.get('title'))
        return (
            0 if candidate_id == active_path_id else 1,
            -family_preference(row.get('family_id')),
            -support_weight(row.get('support_status')),
            -float(row.get('family_score', 0) or 0),
            -float(row.get('confidence_score', 0) or 0),
            -float(row.get('value_score', 0) or 0),
            normalize(row.get('title')).lower(),
        )

    representatives = []
    for group in groups.values():
        representatives.append(sorted(group, key=row_key)[0])
    return sorted(representatives, key=row_key)


def candidate_rows_for_full_evaluation(state, direction_registry, limit=MANUSCRIPT_CANDIDATE_EVALUATION_LIMIT):
    rows = representative_candidate_rows(state, direction_registry)
    if not limit or len(rows) <= limit:
        return rows
    return rows[:limit]


def active_and_watchlist_candidates(candidates):
    active = []
    watchlist = []
    for candidate in candidates:
        status = normalize(candidate.get('publication_status'))
        if status in SUBMISSION_TRACKER_EXIT_STATUSES:
            continue
        if len(active) < ACTIVE_LANE_LIMIT:
            candidate['active_lane'] = True
            candidate['watchlist_rank'] = None
            active.append(candidate)
        else:
            candidate['active_lane'] = False
            candidate['watchlist_rank'] = len(watchlist) + 1
            watchlist.append(candidate)
    return active, watchlist


def queue_summary(active, watchlist, publication_tracker):
    tracker_entries = publication_tracker.get('entries', []) or []
    return {
        'active_count': len(active),
        'watchlist_count': len(watchlist),
        'tracked_publications': len(tracker_entries),
        'ready_for_codex_draft_count': sum(1 for item in active + watchlist if item.get('ready_for_codex_draft')),
        'ready_for_metadata_only_count': sum(1 for item in active + watchlist if item.get('ready_for_metadata_only')),
    }


def build_manuscript_queue_payload(state, direction_registry=None, publication_tracker=None, journal_registry=None):
    direction_registry = direction_registry or {}
    publication_tracker = publication_tracker or load_publication_tracker()
    journal_registry = journal_registry or load_journal_registry()
    indexes = build_phase_indexes(state)
    corpus_index = load_corpus_reference_index()
    connector_context = build_connector_context()
    candidates = []
    publication_entries = {}
    for row in candidate_rows_for_full_evaluation(state, direction_registry):
        candidate, publication_entry = build_candidate_payload(
            row,
            state,
            direction_registry,
            journal_registry,
            publication_tracker,
            indexes,
            corpus_index,
            connector_context=connector_context,
        )
        publication_entries[candidate['candidate_id']] = publication_entry
        candidates.append(candidate)
    candidates.sort(key=lambda item: candidate_priority_key(item, item.get('publication_status')))
    active, watchlist = active_and_watchlist_candidates(candidates)
    tracker_rows = []
    for entry in publication_tracker.get('entries', []) or []:
        candidate_id = canonical_candidate_id(entry.get('candidate_id'))
        candidate = next((item for item in candidates if item['candidate_id'] == candidate_id), None)
        tracker_rows.append({
            'candidate_id': candidate_id,
            'title': normalize(entry.get('title')) or (candidate['title'] if candidate else candidate_id),
            'status': normalize(entry.get('status')),
            'journal_name': normalize(entry.get('journal_name')),
            'updated_at': normalize(entry.get('updated_at')),
            'note': normalize(entry.get('note')),
        })
    draft_outputs = build_generated_draft_summary(active + watchlist)
    return {
        'generated_at': iso_now(),
        'summary': queue_summary(active, watchlist, publication_tracker),
        'connector_status': build_connector_dashboard_payload(connector_context),
        'active_candidates': active,
        'watchlist': watchlist,
        'publication_tracker': tracker_rows,
        'publication_statuses': QUEUE_STATUSES,
        'draft_outputs': draft_outputs,
        'ready_for_metadata_only_candidates': draft_outputs.get('ready_for_metadata_only', []),
    }


def candidate_folder(root, candidate):
    return Path(root) / candidate['pack_key']


def candidate_row_lookup(state, direction_registry=None):
    rows = candidate_rows_for_full_evaluation(state, direction_registry or {})
    lookup = {}
    for row in rows:
        candidate_id = canonical_candidate_id(row.get('candidate_id') or row.get('title'))
        lookup[candidate_id] = row
    return lookup


def materialize_manuscript_outputs(state, direction_registry=None, publication_tracker=None, journal_registry=None, manuscript_root=MANUSCRIPT_ROOT, manuscript_drafts_root=MANUSCRIPT_DRAFTS_ROOT, ready_for_metadata_root=READY_FOR_METADATA_DRAFTS_ROOT, queue_state_path=MANUSCRIPT_QUEUE_STATE_PATH, publication_state_path=PUBLICATION_TRACKER_STATE_PATH):
    direction_registry = direction_registry or {}
    publication_tracker = publication_tracker or load_publication_tracker(publication_state_path)
    journal_registry = journal_registry or load_journal_registry()
    if not Path(publication_state_path).exists():
        write_json(publication_state_path, publication_tracker)
    queue = build_manuscript_queue_payload(state, direction_registry=direction_registry, publication_tracker=publication_tracker, journal_registry=journal_registry)
    row_lookup = candidate_row_lookup(state, direction_registry=direction_registry)
    indexes = build_phase_indexes(state)
    corpus_index = load_corpus_reference_index()
    connector_context = build_connector_context()
    root = Path(manuscript_root)
    root.mkdir(parents=True, exist_ok=True)
    drafts_root = Path(manuscript_drafts_root)
    drafts_root.mkdir(parents=True, exist_ok=True)
    ready_root = Path(ready_for_metadata_root)
    ready_root.mkdir(parents=True, exist_ok=True)

    manifests = []
    draft_manifests = []
    expected_folder_names = set()
    expected_draft_folder_names = set()
    expected_ready_docx_names = set()
    for candidate in queue['active_candidates'] + queue['watchlist']:
        row = row_lookup.get(candidate['candidate_id'])
        if not row:
            continue
        publication_entry = publication_entry_for_candidate(publication_tracker, candidate['candidate_id'])
        folder = candidate_folder(root, candidate)
        draft_folder = candidate_draft_folder(drafts_root, candidate)
        expected_folder_names.add(folder.name)
        expected_draft_folder_names.add(draft_folder.name)
        artifacts_dir = folder / 'artifacts'
        drafts_dir = folder / 'drafts'
        section_packets_dir = folder / 'section_packets'
        task_outputs_dir = folder / 'task_outputs'
        artifacts_dir.mkdir(parents=True, exist_ok=True)
        drafts_dir.mkdir(parents=True, exist_ok=True)
        section_packets_dir.mkdir(parents=True, exist_ok=True)
        task_outputs_dir.mkdir(parents=True, exist_ok=True)
        draft_folder.mkdir(parents=True, exist_ok=True)

        phase_entries = phase_entries_for_candidate(row, state, indexes)
        candidate['connector_summary'] = candidate_connector_summary(candidate, phase_entries, connector_context)
        evidence_bundle = build_evidence_bundle(candidate, row, state, phase_entries, corpus_index)
        candidate['evidence_bundle_summary'] = evidence_bundle['summary']
        candidate['journal_specific_rigor_summary'] = evidence_bundle.get('journal_specific_rigor', {}).get('summary', {})
        candidate['task_ledger'], task_outputs, candidate['task_execution_summary'] = execute_manuscript_tasks(candidate, row, phase_entries, evidence_bundle)
        candidate['top_blocker'] = derive_top_blocker(row, candidate['task_ledger'], evidence_bundle.get('contradictions'))
        candidate['draft_readiness'] = score_draft_readiness(
            row,
            candidate['scientific_strength'],
            candidate['journal_targets'],
            candidate['translation_maturity'],
            tasks=candidate['task_ledger'],
            evidence_bundle_summary=evidence_bundle.get('summary'),
        )
        candidate['draft_readiness_bar'] = score_bar_payload(candidate['draft_readiness'])
        candidate['manuscript_gate_state'] = manuscript_gate_state(candidate['scientific_strength'], candidate['journal_targets'], candidate['draft_readiness'], candidate['translation_maturity'], candidate['task_ledger'])
        candidate['ready_for_codex_draft'] = ready_for_codex_draft(candidate['scientific_strength'], candidate['journal_targets'], candidate['draft_readiness'], candidate['task_ledger'], candidate['manuscript_gate_state'])
        candidate['draft_status'] = 'ready_for_codex_draft' if candidate['ready_for_codex_draft'] else ('pack_ready' if candidate['manuscript_gate_state'] == 'ready to build phase 8 pack' else 'gathering evidence')
        candidate['review_memo_status'] = 'available' if candidate['manuscript_gate_state'] != 'not ready' else 'early warning'
        candidate['recommended_article_type'] = recommended_article_type(candidate)
        candidate['manuscript_draft_title'] = manuscript_draft_title(candidate, row, phase_entries)
        candidate['manuscript_keywords'] = manuscript_keywords(candidate, row, phase_entries)
        missing_packet = hydrate_candidate_draft_output(candidate, evidence_bundle, manuscript_drafts_root=drafts_root)

        manifest = build_pack_manifest(candidate, row, publication_entry)
        manifests.append({'candidate_id': candidate['candidate_id'], 'folder': str(folder.relative_to(REPO_ROOT)), 'gate_state': candidate['manuscript_gate_state']})
        write_json(folder / 'pack_manifest.json', manifest)
        write_json(folder / 'task_ledger.json', {'updated_at': iso_now(), 'summary': candidate['task_execution_summary'], 'tasks': candidate['task_ledger']})
        write_json(folder / 'task_executor_summary.json', candidate['task_execution_summary'])
        write_json(folder / 'journal_targets.json', {
            'updated_at': iso_now(),
            'primary': candidate['journal_targets'].get('primary'),
            'backups': candidate['journal_targets'].get('backups'),
            'scored': candidate['journal_targets'].get('scored'),
        })
        write_text(folder / 'review_memo.md', build_review_memo(candidate, row))
        write_text(folder / 'claim_evidence_map.md', build_claim_evidence_map(candidate, row))
        write_text(folder / 'journal_fit.md', build_journal_fit_markdown(candidate))
        write_text(folder / 'manuscript_brief.md', build_manuscript_brief(candidate, row))
        write_text(folder / 'codex_handoff.md', build_codex_handoff_markdown(candidate))
        if 'resolve_lead_contradiction' in task_outputs:
            write_text(folder / 'blocking_resolution_plan.md', task_outputs['resolve_lead_contradiction'])
        if 'strengthen_translational_packet' in task_outputs:
            write_text(folder / 'translational_strengthening_plan.md', task_outputs['strengthen_translational_packet'])
        write_json(folder / 'references.json', {
            'updated_at': iso_now(),
            'summary': evidence_bundle['summary'],
            'references': evidence_bundle['references'],
        })
        write_text(folder / 'references.md', evidence_bundle['reference_markdown'])
        write_json(folder / 'claim_support_matrix.json', {
            'updated_at': iso_now(),
            'claims': evidence_bundle['claims'],
        })
        write_text(folder / 'claim_support_matrix.md', evidence_bundle['claim_support_markdown'])
        write_json(folder / 'contradiction_register.json', {
            'updated_at': iso_now(),
            'issues': evidence_bundle['contradictions'],
        })
        write_text(folder / 'contradiction_register.md', evidence_bundle['contradiction_markdown'])
        write_json(folder / 'figure_evidence_plan.json', {
            'updated_at': iso_now(),
            'figures': evidence_bundle['figures'],
        })
        write_text(folder / 'figure_evidence_plan.md', evidence_bundle['figure_plan_markdown'])
        write_json(folder / 'source_artifact_manifest.json', {
            'updated_at': iso_now(),
            'artifacts': evidence_bundle['source_artifacts'],
        })
        write_text(folder / 'source_artifact_manifest.md', evidence_bundle['source_artifact_markdown'])
        write_json(folder / 'data_points.json', {
            'updated_at': iso_now(),
            'data_points': evidence_bundle['data_points'],
        })
        write_text(folder / 'data_points.md', evidence_bundle['data_points_markdown'])
        write_json(folder / 'journal_requirements_snapshot.json', {
            'updated_at': iso_now(),
            'journal_snapshot': evidence_bundle['journal_snapshot'],
        })
        write_json(folder / 'journal_specific_rigor_checklist.json', {
            'updated_at': iso_now(),
            'journal_specific_rigor': evidence_bundle['journal_specific_rigor'],
        })
        write_text(folder / 'journal_specific_rigor_checklist.md', evidence_bundle['journal_specific_rigor_markdown'])
        write_json(folder / 'journal_specific_rigor_fill_packet.json', {
            'updated_at': iso_now(),
            'journal_specific_rigor_fill_packet': evidence_bundle['journal_specific_rigor_fill_packet'],
        })
        write_text(folder / 'journal_specific_rigor_fill_packet.md', evidence_bundle['journal_specific_rigor_fill_markdown'])
        write_json(folder / 'evidence_bundle_summary.json', {
            'updated_at': iso_now(),
            'summary': evidence_bundle['summary'],
        })
        for filename, content in evidence_bundle['section_packets'].items():
            write_text(section_packets_dir / filename, content)
        expected_task_output_names = {f'{task_id}.md' for task_id in task_outputs}
        for existing_output in task_outputs_dir.glob('*.md'):
            if existing_output.name not in expected_task_output_names:
                existing_output.unlink()
        for task_id, content in task_outputs.items():
            write_text(task_outputs_dir / f'{task_id}.md', content)
        write_json(folder / 'candidate_snapshot.json', candidate)
        ready_path = folder / READY_FOR_CODEX_FILENAME
        if candidate['ready_for_codex_draft']:
            write_json(ready_path, {
                'updated_at': iso_now(),
                'candidate_id': candidate['candidate_id'],
                'title': candidate['title'],
                'primary_journal': candidate['journal_targets'].get('primary'),
                'backups': candidate['journal_targets'].get('backups'),
                'drafting_mode': 'journal_aware_markdown_first',
                'required_files': [
                    'pack_manifest.json',
                    'references.json',
                    'claim_support_matrix.json',
                    'contradiction_register.json',
                    'source_artifact_manifest.json',
                    'data_points.json',
                    'journal_requirements_snapshot.json',
                    'journal_specific_rigor_checklist.json',
                    'journal_specific_rigor_checklist.md',
                    'journal_specific_rigor_fill_packet.json',
                    'journal_specific_rigor_fill_packet.md',
                    'blocking_resolution_plan.md',
                    'translational_strengthening_plan.md',
                    'review_memo.md',
                    'journal_fit.md',
                    'codex_handoff.md',
                    'section_packets/00_introduction_packet.md',
                    'section_packets/01_methods_packet.md',
                    'section_packets/02_results_packet.md',
                    'section_packets/03_discussion_packet.md',
                    'section_packets/04_limitations_packet.md',
                ],
            })
        elif ready_path.exists():
            ready_path.unlink()

        write_text(draft_folder / 'draft_status.md', build_draft_status_markdown(candidate, row, phase_entries, missing_packet))
        write_text(draft_folder / 'target_journal.md', build_target_journal_markdown(candidate))
        write_text(draft_folder / 'comprehensive_outline.md', build_comprehensive_outline_markdown(candidate, row, phase_entries, evidence_bundle, missing_packet))
        write_text(draft_folder / 'missing_metadata_and_items.md', build_missing_metadata_markdown(candidate, missing_packet))
        manuscript_markdown = build_full_manuscript_draft_markdown(candidate, row, phase_entries, evidence_bundle, missing_packet)
        write_text(draft_folder / 'manuscript_draft.md', manuscript_markdown)
        draft_manifest = {
            'updated_at': iso_now(),
            'candidate_id': candidate['candidate_id'],
            'pack_key': candidate['pack_key'],
            'title': candidate['title'],
            'draft_title': candidate.get('manuscript_draft_title'),
            'generated_draft_status': candidate.get('generated_draft_status'),
            'ready_for_metadata_only': bool(candidate.get('ready_for_metadata_only')),
            'primary_journal': candidate['journal_targets'].get('primary'),
            'recommended_article_type': candidate.get('recommended_article_type'),
            'draft_output': candidate.get('draft_output', {}),
            'missing_metadata_summary': candidate.get('missing_metadata_summary', {}),
            'evidence_bundle_summary': evidence_bundle.get('summary', {}),
            'manuscript_gate_state': candidate.get('manuscript_gate_state'),
        }
        write_json(draft_folder / GENERATED_DRAFT_MANIFEST_FILENAME, draft_manifest)
        draft_manifests.append(draft_manifest)

        ready_docx_path = Path(candidate.get('draft_output', {}).get('ready_metadata_docx_absolute_path') or ready_for_metadata_docx_path(ready_root, candidate))
        if candidate.get('ready_for_metadata_only'):
            write_ready_metadata_docx(manuscript_markdown, ready_docx_path, candidate)
            expected_ready_docx_names.add(ready_docx_path.name)
        elif ready_docx_path.exists():
            ready_docx_path.unlink()

    prune_stale_outputs = os.environ.get('PHASE8_PRUNE_STALE_OUTPUTS', '').strip() == '1'
    if prune_stale_outputs:
        for child in root.iterdir():
            if not child.is_dir() or child.name in expected_folder_names:
                continue
            if (child / 'pack_manifest.json').exists() and (child / 'candidate_snapshot.json').exists():
                try:
                    shutil.rmtree(child)
                except FileNotFoundError:
                    # Another rebuild pass may have already removed the stale folder.
                    pass

        for child in drafts_root.iterdir():
            if not child.is_dir() or child.name in expected_draft_folder_names:
                continue
            if (child / GENERATED_DRAFT_MANIFEST_FILENAME).exists():
                try:
                    shutil.rmtree(child)
                except FileNotFoundError:
                    pass

    for existing_docx in ready_root.glob('*.docx'):
        if existing_docx.name not in expected_ready_docx_names:
            try:
                existing_docx.unlink()
            except FileNotFoundError:
                pass

    queue['draft_outputs'] = build_generated_draft_summary(queue['active_candidates'] + queue['watchlist'], root=drafts_root, ready_root=ready_root)
    queue['ready_for_metadata_only_candidates'] = [
        candidate
        for candidate in queue['active_candidates'] + queue['watchlist']
        if candidate.get('ready_for_metadata_only')
    ]
    queue['artifact_manifest'] = manifests
    write_json(drafts_root / GENERATED_DRAFT_INDEX_FILENAME, {
        'updated_at': iso_now(),
        'summary': queue.get('summary', {}),
        'draft_outputs': queue.get('draft_outputs', {}),
        'draft_manifests': draft_manifests,
    })
    write_json(STATE_DIR / READY_FOR_METADATA_INDEX_FILENAME, {
        'updated_at': iso_now(),
        'ready_for_metadata_root_relative_path': relative_path_or_blank(ready_root),
        'ready_for_metadata_root_absolute_path': absolute_path_or_blank(ready_root),
        'docx_files': [
            {
                'candidate_id': candidate.get('candidate_id'),
                'title': candidate.get('title'),
                'docx_relative_path': candidate.get('draft_output', {}).get('ready_metadata_docx_relative_path', ''),
                'docx_absolute_path': candidate.get('draft_output', {}).get('ready_metadata_docx_absolute_path', ''),
            }
            for candidate in queue['ready_for_metadata_only_candidates']
        ],
    })
    write_json(queue_state_path, queue)
    return queue
